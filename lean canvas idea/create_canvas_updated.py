from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Lean Business Model Canvas', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('GreenGrid AI')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(26, 95, 42)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle2.add_run('An AI-Powered Energy Optimization Platform for Government and Public Sector Infrastructure')
run2.italic = True
run2.font.size = Pt(12)

doc.add_paragraph()

# Helper function to set cell shading
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Create the Lean Canvas Table (Visual Single-Page Format)
doc.add_heading('LEAN CANVAS - SINGLE PAGE VIEW', level=1)

# Create main canvas table
canvas = doc.add_table(rows=3, cols=5)
canvas.style = 'Table Grid'

# Define colors
colors = {
    'problem': 'FFE5E5',      # Light red
    'solution': 'E5F0FF',     # Light blue
    'metrics': 'FFF5E5',      # Light orange
    'uvp': 'F5E5FF',          # Light purple
    'advantage': 'FFF9E5',    # Light yellow
    'channels': 'E5FFF0',     # Light teal
    'segments': 'E5FFFA',     # Light cyan
    'costs': 'FFEBE5',        # Light coral
    'revenue': 'E5FFE5'       # Light green
}

# Row 0-1, Col 0: Problem
cell = canvas.cell(0, 0).merge(canvas.cell(1, 0))
set_cell_shading(cell, colors['problem'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('PROBLEM\n').bold = True
p.runs[0].font.color.rgb = RGBColor(231, 76, 60)
p.add_run('\n1. Energy Waste: 20-30% in public buildings\n\n2. Sustainability Mandates: Net-zero targets without data\n\n3. Budget Constraints: Rising costs, limited expertise\n\n').font.size = Pt(9)
p.add_run('Existing Alternatives:\n').bold = True
p.add_run('- Manual audits\n- Basic BMS\n- Spreadsheets').font.size = Pt(8)

# Row 0, Col 1: Solution
cell = canvas.cell(0, 1)
set_cell_shading(cell, colors['solution'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('SOLUTION\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(52, 152, 219)
p.add_run('- Predictive Energy Analytics (ML forecasting)\n- Automated Optimization Engine\n- Sustainability Dashboard\n- Renewable Integration').font.size = Pt(9)

# Row 1, Col 1: Key Metrics
cell = canvas.cell(1, 1)
set_cell_shading(cell, colors['metrics'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('KEY METRICS\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(230, 126, 34)
p.add_run('- Energy Savings (%)\n- Cost Reduction ($)\n- Carbon Reduction (tCO2)\n- Facilities Onboarded\n- MRR Growth\n- Prediction Accuracy').font.size = Pt(9)

# Row 0-1, Col 2: UVP
cell = canvas.cell(0, 2).merge(canvas.cell(1, 2))
set_cell_shading(cell, colors['uvp'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('UNIQUE VALUE PROPOSITION\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(155, 89, 182)
p.add_run('"Transform public buildings into smart, sustainable assets. 15-25% energy savings through predictive intelligence."\n\n').font.size = Pt(9)
p.runs[-1].italic = True
p.add_run('High-Level Concept:\n').bold = True
p.add_run('"Nest for government infrastructure"\n\n').font.size = Pt(9)
p.add_run('Differentiators:\n').bold = True
p.add_run('- Public sector purpose-built\n- No hardware required\n- Transparent AI\n- ROI in 6-12 months').font.size = Pt(9)

# Row 0, Col 3: Unfair Advantage
cell = canvas.cell(0, 3)
set_cell_shading(cell, colors['advantage'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('UNFAIR ADVANTAGE\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(243, 156, 18)
p.add_run('- Proprietary govt energy dataset\n- Public sector expertise\n- Strategic utility partnerships\n- Regulatory head start').font.size = Pt(9)

# Row 1, Col 3: Channels
cell = canvas.cell(1, 3)
set_cell_shading(cell, colors['channels'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('CHANNELS\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(22, 160, 133)
p.add_run('- Direct govt sales\n- Govt marketplaces\n- Smart city conferences\n- Partner networks\n- 90-day pilots').font.size = Pt(9)

# Row 0-1, Col 4: Customer Segments
cell = canvas.cell(0, 4).merge(canvas.cell(1, 4))
set_cell_shading(cell, colors['segments'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('CUSTOMER SEGMENTS\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(26, 188, 156)
p.add_run('Primary:\n').bold = True
p.add_run('- Municipal governments\n- School districts\n- Public hospitals\n\n').font.size = Pt(9)
p.add_run('Secondary:\n').bold = True
p.add_run('- State agencies\n- Universities\n\n').font.size = Pt(9)
p.add_run('Early Adopters:\n').bold = True
p.add_run('- Climate emergency cities\n- Smart city initiatives').font.size = Pt(9)

# Row 2, Col 0-1: Cost Structure
cell = canvas.cell(2, 0).merge(canvas.cell(2, 1))
set_cell_shading(cell, colors['costs'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('COST STRUCTURE\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(192, 57, 43)
p.add_run('Fixed: ').bold = True
p.add_run('Cloud (AWS/Azure Gov), Engineering, Compliance\n').font.size = Pt(9)
p.add_run('Variable: ').bold = True
p.add_run('Sales, Support, Marketing\n').font.size = Pt(9)
p.add_run('Split: ').bold = True
p.add_run('Eng 40% | Sales 25% | Infra 15% | Support 12% | G&A 8%').font.size = Pt(9)

# Row 2, Col 2-4: Revenue Streams
cell = canvas.cell(2, 2).merge(canvas.cell(2, 4))
set_cell_shading(cell, colors['revenue'])
cell.text = ''
p = cell.paragraphs[0]
p.add_run('REVENUE STREAMS\n\n').bold = True
p.runs[0].font.color.rgb = RGBColor(39, 174, 96)
p.add_run('SaaS Tiers: ').bold = True
p.add_run('Starter ($500-1.5K/mo) | Professional ($3-10K/mo) | Enterprise ($50-500K/yr)\n').font.size = Pt(9)
p.add_run('Secondary: ').bold = True
p.add_run('Implementation, Training, Premium Support\n').font.size = Pt(9)
p.add_run('Advantages: ').bold = True
p.add_run('Recurring revenue, Multi-year contracts, Expansion revenue').font.size = Pt(9)

doc.add_paragraph()

# Page break for Market Validation
doc.add_page_break()

# Market Validation Summary with new interviews
doc.add_heading('MARKET VALIDATION SUMMARY', level=1)

doc.add_paragraph().add_run('Research Methodology').bold = True
doc.add_paragraph('To validate the GreenGrid AI business concept, primary research was conducted through interviews with potential users and domain experts in the public sector energy management and sustainability space. This research aimed to understand current pain points, existing solutions, and receptivity to AI-powered alternatives in the Irish and broader European context.')

doc.add_paragraph()
doc.add_paragraph().add_run('Interview Participants').bold = True

# Participant 1 - Potential User
p = doc.add_paragraph()
p.add_run('Participant A - Sarah M., Facilities Coordinator, Dublin City Council').bold = True
doc.add_paragraph('Sarah oversees energy management for 8 public buildings across Dublin, including civic offices, libraries, and community centres. With 7 years of experience in facilities management, she is responsible for implementing the Council\'s Sustainable Energy Action Plan and reporting on building performance to senior management. She expressed significant frustration with current manual processes.')

doc.add_paragraph()

# Participant 2 - Domain Expert (Irish Sustainability Firm)
p = doc.add_paragraph()
p.add_run('Participant B - Cian O\'Sullivan, Senior Sustainability Consultant, Greenway Solutions Ltd. (Cork)').bold = True
doc.add_paragraph('Cian works at Greenway Solutions, an Irish sustainability consultancy firm based in Cork that specialises in helping public and private sector organisations achieve their decarbonisation targets. With a background in environmental engineering from UCC and 5 years of experience advising local authorities across Munster, Cian provides strategic guidance on energy efficiency projects and ESG compliance. He has worked on over 30 public sector sustainability projects across Ireland.')

doc.add_paragraph()
doc.add_paragraph().add_run('Key Insights from Interviews').bold = True

# Interview insights
insights = [
    ('Pain Point Validation',
     'Both participants confirmed that energy waste in public buildings is a significant and ongoing challenge in Ireland. Sarah noted: "We\'re still relying on quarterly energy bills to understand our consumption. By the time I spot an issue, we\'ve already wasted thousands of euros. The council has ambitious climate targets for 2030, but we don\'t have the real-time visibility to track whether we\'re on course." Cian added from his consultancy perspective: "I see this across every local authority I work with. They have the political will to decarbonise, but the tools haven\'t caught up with the ambition."'),

    ('Current Solutions Are Inadequate',
     'Existing alternatives were described as fragmented and insufficient. Sarah explained: "We have a basic BMS in our main civic building, but it\'s essentially a glorified thermostat. No predictive capability, no integration with weather data, nothing automated. For our smaller buildings, we\'re literally tracking energy use in Excel spreadsheets." Cian confirmed this pattern: "The standard approach in Ireland is still annual energy audits. You get a nice PDF report, implement a few recommendations, and then nothing until next year. It\'s completely reactive. What councils need is continuous intelligence."'),

    ('AI Receptivity and Irish Market Context',
     'There was strong interest in AI-powered solutions, with both participants noting increased openness in the Irish public sector. Cian observed: "The Climate Action Plan 2024 has put real pressure on local authorities. They\'re now mandated to reduce emissions by 51% by 2030. That regulatory push is making councils much more receptive to innovative solutions. Three years ago, AI in building management would have seemed futuristic\u2014now it\'s becoming necessary." Sarah agreed: "If you can show me a system that integrates with our existing meters and gives me actionable insights without a massive IT project, I\'m interested. We don\'t have budget for hardware, but we have budget for solutions that demonstrably cut costs."'),

    ('Integration and Data Security Concerns',
     'Both participants emphasized the importance of working with existing infrastructure. Sarah stated: "Any solution has to work with what we\'ve got. We\'ve smart meters in most buildings now thanks to SEAI grants, but they\'re not being used to their full potential. I need something that pulls that data together and tells me what to do with it." Cian raised the public sector procurement angle: "GDPR compliance and data residency are non-negotiable for Irish local authorities. If you\'re storing building energy data, it needs to be clear where that data lives and who has access. Government clients will ask these questions upfront."'),

    ('Willingness to Pay and ROI Expectations',
     'Both participants indicated budget availability for solutions with demonstrable ROI. Sarah confirmed: "My annual energy budget is roughly \u20ac400,000 across our buildings. If you can prove 15-20% savings, that\'s \u20ac60-80K back into public services. The business case practically writes itself, and it helps us meet our public sector energy efficiency targets under SI 426." Cian added: "I\'ve seen councils approve \u20ac50-100K annual contracts for energy services that deliver measurable results. The key is proving the savings are real and sustainable. A 90-day pilot with verified metrics would be exactly what procurement teams need to justify the spend."')
]

for title, insight in insights:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(insight)
    doc.add_paragraph()

doc.add_paragraph().add_run('Conclusion').bold = True
conclusion = doc.add_paragraph()
conclusion.add_run('The market validation research strongly supports the GreenGrid AI value proposition within the Irish and broader European public sector context. The identified problems are real and pressing, driven by regulatory mandates like Ireland\'s Climate Action Plan and the EU\'s Energy Efficiency Directive. Current solutions are insufficient, relying on manual processes and periodic audits rather than continuous intelligence. Both the end-user (Sarah) and domain expert (Cian) perspectives confirm clear demand for an AI-powered platform that delivers measurable energy savings without requiring significant hardware investment. The insights gathered have directly shaped the product features, particularly the emphasis on integration with existing smart meter infrastructure, GDPR-compliant data handling, transparent AI recommendations, and ROI-focused reporting aligned with Irish public sector energy reporting requirements (SI 426).')

doc.add_paragraph()

# Additional section for AI/Data Science value
doc.add_page_break()
doc.add_heading('HOW AI/DATA SCIENCE CREATES VALUE', level=1)

value_table = doc.add_table(rows=7, cols=3)
value_table.style = 'Table Grid'

value_data = [
    ('Value Driver', 'AI/Data Science Contribution', 'Business Impact'),
    ('Cost Reduction', 'Predictive models optimize energy consumption', '15-25% reduction in energy bills'),
    ('Risk Mitigation', 'Anomaly detection prevents equipment failures', 'Reduced downtime and emergency repairs'),
    ('Compliance', 'Automated carbon tracking and reporting', 'Meets SI 426 and EU requirements'),
    ('Decision Support', 'Data-driven recommendations for investments', 'Better allocation of public funds'),
    ('Scalability', 'ML models improve with more data', 'Network effects benefit all customers'),
    ('Speed', 'Real-time processing vs. manual audits', 'Immediate insights, not quarterly reports')
]

for i, (driver, contrib, impact) in enumerate(value_data):
    row = value_table.rows[i]
    row.cells[0].text = driver
    row.cells[1].text = contrib
    row.cells[2].text = impact
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E5FFE5')

doc.add_paragraph()

doc.add_paragraph().add_run('The Data Flywheel Effect:').bold = True
flywheel = doc.add_paragraph()
flywheel.add_run('More Government Customers \u2192 More Building Data \u2192 Better ML Accuracy \u2192 Greater Savings \u2192 Stronger Case Studies \u2192 Easier Acquisition \u2192 (Cycle Repeats)')

doc.add_paragraph()

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Document prepared for: CSC1112 Domain Applications 3 - Lean Canvas Assignment').italic = True
footer.add_run('\n')
footer.add_run('Dublin City University | April 2026').italic = True

# Save
doc.save(r'D:\Final Year\App Domains 3\lean canvas idea\Lean_Business_Model_Canvas_Updated.docx')
print('Word document created successfully!')
