from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Lean Business Model Canvas', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('GreenGrid AI')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 100, 0)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle2.add_run('An AI-Powered Energy Optimization Platform for Government and Public Sector Infrastructure')
run2.italic = True
run2.font.size = Pt(12)

doc.add_paragraph()

# Section 1: Problem
doc.add_heading('1. PROBLEM', level=1)

doc.add_heading('Top 3 Problems Addressed:', level=2)

p1 = doc.add_paragraph()
p1.add_run('1. Inefficient Energy Consumption in Public Buildings').bold = True
doc.add_paragraph('Government buildings (offices, schools, hospitals, libraries) often waste 20-30% of energy due to outdated HVAC systems, poor scheduling, and lack of real-time monitoring. Manual energy management is reactive rather than predictive.', style='List Bullet')

p2 = doc.add_paragraph()
p2.add_run('2. Difficulty Meeting Sustainability Mandates').bold = True
doc.add_paragraph('Governments face increasing pressure to meet net-zero targets and climate commitments (e.g., Paris Agreement, local climate action plans). There is a lack of actionable data to track progress toward emissions reduction goals.', style='List Bullet')

p3 = doc.add_paragraph()
p3.add_run('3. High Operational Costs and Budget Constraints').bold = True
doc.add_paragraph('Public sector organizations face tight budgets while energy costs continue to rise. Limited technical expertise exists to implement and manage complex energy systems.', style='List Bullet')

doc.add_heading('Existing Alternatives:', level=2)
alternatives = ['Manual energy audits (expensive, infrequent, static reports)',
                'Basic Building Management Systems (BMS) without predictive capabilities',
                'Spreadsheet-based tracking and reporting',
                'External consultants (costly, not real-time)']
for alt in alternatives:
    doc.add_paragraph(alt, style='List Bullet')

# Section 2: Solution
doc.add_heading('2. SOLUTION', level=1)
doc.add_heading('AI-Powered Energy Intelligence Platform', level=2)

doc.add_paragraph().add_run('Core Features:').bold = True

solutions = [
    ('Predictive Energy Analytics', 'Machine learning models that forecast energy demand based on weather, occupancy patterns, historical usage, and seasonal trends. Anomaly detection to identify equipment failures or energy waste in real-time.'),
    ('Automated Optimization Engine', 'AI-driven recommendations for HVAC, lighting, and equipment scheduling. Integration with smart building systems for automated load balancing. Peak demand management to reduce utility costs.'),
    ('Sustainability Dashboard & Reporting', 'Real-time carbon footprint tracking across all facilities. Automated compliance reporting for government sustainability mandates. Progress visualization toward net-zero goals.'),
    ('Renewable Energy Integration', 'Forecasting for solar/wind generation on government properties. Optimal battery storage management. Grid interaction optimization (when to buy/sell energy).')
]

for title, desc in solutions:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# Section 3: Key Metrics
doc.add_heading('3. KEY METRICS', level=1)

metrics_table = doc.add_table(rows=9, cols=2)
metrics_table.style = 'Table Grid'

metrics_data = [
    ('Metric', 'Description'),
    ('Energy Savings (%)', 'Reduction in kWh consumption per facility'),
    ('Cost Savings ($)', 'Monthly/annual reduction in energy bills'),
    ('Carbon Reduction (tons CO2)', 'Measurable emissions avoided'),
    ('Facilities Onboarded', 'Number of buildings using the platform'),
    ('Monthly Recurring Revenue (MRR)', 'Subscription revenue growth'),
    ('Customer Retention Rate', 'Annual contract renewals'),
    ('Platform Uptime', 'System reliability (target: 99.9%)'),
    ('Prediction Accuracy', 'ML model performance (energy forecast accuracy)')
]

for i, (metric, desc) in enumerate(metrics_data):
    row = metrics_table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = desc
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 4: Unique Value Proposition
doc.add_heading('4. UNIQUE VALUE PROPOSITION', level=1)

doc.add_paragraph().add_run('Single, Clear, Compelling Message:').bold = True

quote = doc.add_paragraph()
quote.paragraph_format.left_indent = Inches(0.5)
quote.paragraph_format.right_indent = Inches(0.5)
run = quote.add_run('"Transform public buildings into smart, sustainable assets. GreenGrid AI delivers 15-25% energy savings through predictive intelligence—helping governments meet climate goals while reducing operational costs."')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('High-Level Concept: ').bold = True
p.add_run('"Nest thermostat intelligence, scaled for government infrastructure"')

doc.add_paragraph().add_run('Key Differentiators:').bold = True
differentiators = [
    'Purpose-built for public sector compliance and procurement requirements',
    'No hardware required — integrates with existing smart meters and BMS',
    'Transparent AI — explainable recommendations that satisfy audit requirements',
    'Proven ROI — typically pays for itself within 6-12 months'
]
for d in differentiators:
    doc.add_paragraph(d, style='List Bullet')

# Section 5: Unfair Advantage
doc.add_heading('5. UNFAIR ADVANTAGE', level=1)
doc.add_paragraph().add_run('What Cannot Be Easily Copied or Bought:').bold = True

advantages = [
    ('Proprietary Government Energy Dataset', 'Aggregated, anonymized data from public sector buildings creates increasingly accurate models. Network effects: more facilities = better predictions for all customers.'),
    ('Domain Expertise in Public Sector Procurement', 'Deep understanding of government buying cycles, compliance requirements, and stakeholder management. Pre-approved vendor status and security certifications (e.g., FedRAMP, ISO 27001).'),
    ('Strategic Partnerships', 'Exclusive integrations with major utility companies and smart meter providers. Relationships with government sustainability offices and policy advisors.'),
    ('Regulatory Head Start', 'Early mover in meeting emerging mandatory ESG reporting requirements. Platform designed around anticipated regulations.')
]

for i, (title, desc) in enumerate(advantages, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {title}').bold = True
    doc.add_paragraph(desc, style='List Bullet')

# Section 6: Customer Segments
doc.add_heading('6. CUSTOMER SEGMENTS', level=1)

doc.add_paragraph().add_run('Primary Segment: Municipal & Local Governments').bold = True
primary = ['City councils and local authorities', 'Public school districts', 'Municipal utilities', 'Public hospitals and healthcare facilities']
for p in primary:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph().add_run('Secondary Segment: State/Provincial Governments').bold = True
secondary = ['State-owned buildings and facilities', 'Transportation authorities', 'University systems']
for s in secondary:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph().add_run('Tertiary Segment: Federal/National Agencies').bold = True
tertiary = ['Government office buildings', 'Military bases (non-classified facilities)', 'National parks and public lands facilities']
for t in tertiary:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph().add_run('Early Adopters:').bold = True
adopters = ['Progressive cities with declared climate emergencies', 'Municipalities with existing smart city initiatives', 'Government bodies facing public pressure on sustainability', 'Facilities managers seeking to modernize aging infrastructure']
for a in adopters:
    doc.add_paragraph(a, style='List Bullet')

# Section 7: Channels
doc.add_heading('7. CHANNELS', level=1)

channels_table = doc.add_table(rows=8, cols=2)
channels_table.style = 'Table Grid'

channels_data = [
    ('Channel', 'Strategy'),
    ('Direct Sales', 'Dedicated government sales team experienced in public procurement'),
    ('Government Marketplaces', 'Listed on GSA Schedule, G-Cloud, and regional procurement portals'),
    ('Industry Conferences', 'Presence at Smart Cities Expo, GovTech Summit, APPA conferences'),
    ('Partnerships', 'Channel partnerships with energy consultants and system integrators'),
    ('Pilot Programs', 'Free 90-day pilots to demonstrate ROI before full deployment'),
    ('Case Studies', 'Published success stories with measurable outcomes'),
    ('Referrals', 'Government-to-government recommendations through sustainability networks')
]

for i, (channel, strategy) in enumerate(channels_data):
    row = channels_table.rows[i]
    row.cells[0].text = channel
    row.cells[1].text = strategy
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 8: Cost Structure
doc.add_heading('8. COST STRUCTURE', level=1)

doc.add_paragraph().add_run('Fixed Costs:').bold = True
fixed = ['Cloud infrastructure (AWS/Azure Government Cloud)', 'Engineering team salaries (ML engineers, data scientists, developers)', 'Security and compliance certifications', 'Office and administrative overhead']
for f in fixed:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph().add_run('Variable Costs:').bold = True
variable = ['Customer success and support staff', 'Sales commissions', 'Data acquisition and API costs', 'Marketing and conference attendance']
for v in variable:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph().add_run('Estimated Cost Breakdown:').bold = True

cost_table = doc.add_table(rows=6, cols=2)
cost_table.style = 'Table Grid'

cost_data = [
    ('Category', '% of Budget'),
    ('Engineering & Product', '40%'),
    ('Sales & Marketing', '25%'),
    ('Cloud Infrastructure', '15%'),
    ('Customer Success', '12%'),
    ('G&A / Compliance', '8%')
]

for i, (cat, pct) in enumerate(cost_data):
    row = cost_table.rows[i]
    row.cells[0].text = cat
    row.cells[1].text = pct
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 9: Revenue Streams
doc.add_heading('9. REVENUE STREAMS', level=1)

doc.add_paragraph().add_run('Primary: SaaS Subscription (Tiered Pricing)').bold = True

revenue_table = doc.add_table(rows=4, cols=3)
revenue_table.style = 'Table Grid'

revenue_data = [
    ('Tier', 'Target', 'Pricing Model'),
    ('Starter', 'Single facility (<50,000 sq ft)', '$500-1,500/month'),
    ('Professional', 'Multiple facilities (campus/district)', '$3,000-10,000/month'),
    ('Enterprise', 'City-wide or state deployment', 'Custom pricing ($50K-500K/year)')
]

for i, (tier, target, price) in enumerate(revenue_data):
    row = revenue_table.rows[i]
    row.cells[0].text = tier
    row.cells[1].text = target
    row.cells[2].text = price
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_paragraph().add_run('Secondary Revenue:').bold = True
secondary_rev = ['Implementation Services: One-time setup and integration fees', 'Training & Certification: Workshops for facilities managers', 'Premium Support: Dedicated account management and SLA guarantees', 'Data Insights Reports: Custom analytics and benchmarking reports']
for s in secondary_rev:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph().add_run('Revenue Model Advantages:').bold = True
rev_adv = ['Recurring revenue provides predictable cash flow', 'Government contracts typically multi-year (3-5 years)', 'Expansion revenue as customers add facilities']
for r in rev_adv:
    doc.add_paragraph(r, style='List Bullet')

# Section 10: Data Assets & AI Capabilities
doc.add_heading('10. DATA ASSETS & AI CAPABILITIES', level=1)

doc.add_paragraph().add_run('Data Sources:').bold = True

data_table = doc.add_table(rows=7, cols=3)
data_table.style = 'Table Grid'

data_sources = [
    ('Data Type', 'Source', 'Use Case'),
    ('Energy Consumption', 'Smart meters, utility APIs', 'Baseline analysis, anomaly detection'),
    ('Weather Data', 'National weather services, APIs', 'Demand forecasting'),
    ('Occupancy Data', 'Building sensors, schedules', 'HVAC optimization'),
    ('Equipment Telemetry', 'IoT sensors, BMS integration', 'Predictive maintenance'),
    ('Grid Pricing', 'Utility rate schedules', 'Cost optimization'),
    ('Solar/Wind Output', 'On-site renewable sensors', 'Generation forecasting')
]

for i, (dtype, source, use) in enumerate(data_sources):
    row = data_table.rows[i]
    row.cells[0].text = dtype
    row.cells[1].text = source
    row.cells[2].text = use
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_paragraph().add_run('AI/ML Models Deployed:').bold = True

models = [
    ('Energy Demand Forecasting', 'LSTM neural networks, XGBoost ensemble', 'Predict hourly/daily energy consumption', '>95% for 24-hour forecasts'),
    ('Anomaly Detection', 'Isolation Forest, Autoencoders', 'Identify equipment faults, unusual consumption patterns', 'Early warning system for maintenance'),
    ('Optimization Engine', 'Reinforcement Learning, Linear Programming', 'Real-time HVAC and load scheduling', 'Automated control signals or human recommendations'),
    ('Carbon Calculation Model', 'Emissions factor database + consumption data', 'Accurate Scope 1 & 2 emissions tracking', 'Aligned with GHG Protocol standards')
]

for i, (name, algo, purpose, output) in enumerate(models, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. {name}').bold = True
    doc.add_paragraph(f'Algorithm: {algo}', style='List Bullet')
    doc.add_paragraph(f'Purpose: {purpose}', style='List Bullet')
    doc.add_paragraph(f'Output/Accuracy: {output}', style='List Bullet')

doc.add_paragraph().add_run('Data Governance:').bold = True
governance = ['All data encrypted at rest and in transit', 'Government-grade security (SOC 2 Type II, FedRAMP where applicable)', 'Data anonymization for cross-customer model training', 'Customer owns their data; full export capability']
for g in governance:
    doc.add_paragraph(g, style='List Bullet')

# Section 11: How AI/Data Science Creates Value
doc.add_heading('11. HOW AI/DATA SCIENCE CREATES VALUE', level=1)

value_table = doc.add_table(rows=7, cols=3)
value_table.style = 'Table Grid'

value_data = [
    ('Value Driver', 'AI/Data Science Contribution', 'Business Impact'),
    ('Cost Reduction', 'Predictive models optimize energy consumption', '15-25% reduction in energy bills'),
    ('Risk Mitigation', 'Anomaly detection prevents equipment failures', 'Reduced downtime and emergency repairs'),
    ('Compliance', 'Automated carbon tracking and reporting', 'Meets regulatory requirements efficiently'),
    ('Decision Support', 'Data-driven recommendations for capital investments', 'Better allocation of public funds'),
    ('Scalability', 'ML models improve with more data', 'Each new customer improves the product'),
    ('Speed', 'Real-time processing vs. manual audits', 'Immediate insights instead of quarterly reports')
]

for i, (driver, contrib, impact) in enumerate(value_data):
    row = value_table.rows[i]
    row.cells[0].text = driver
    row.cells[1].text = contrib
    row.cells[2].text = impact
    if i == 0:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_paragraph().add_run('The Data Flywheel Effect:').bold = True
flywheel = doc.add_paragraph()
flywheel.add_run('More Government Customers → More Building Data Collected → Better ML Model Accuracy → Greater Energy Savings Delivered → Stronger ROI & Case Studies → Easier Customer Acquisition → (Cycle Repeats)')

# Section 12: Competitive Landscape
doc.add_heading('12. COMPETITIVE LANDSCAPE', level=1)

comp_table = doc.add_table(rows=5, cols=4)
comp_table.style = 'Table Grid'

comp_data = [
    ('Competitor', 'Strengths', 'Weaknesses', 'Our Differentiation'),
    ('Siemens/Honeywell', 'Established, hardware + software', 'Expensive, complex, hardware-dependent', 'Software-only, faster deployment'),
    ('Engie/Schneider', 'Full energy services', 'Generalist, not AI-first', 'AI-native, specialized models'),
    ('Startup competitors', 'Agile, modern tech', 'Limited govt experience', 'Purpose-built for public sector'),
    ('In-house solutions', 'Customized', 'Resource-intensive, no network effects', 'Proven platform, shared learnings')
]

for i, (comp, strength, weak, diff) in enumerate(comp_data):
    row = comp_table.rows[i]
    row.cells[0].text = comp
    row.cells[1].text = strength
    row.cells[2].text = weak
    row.cells[3].text = diff
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 13: Go-To-Market Strategy
doc.add_heading('13. GO-TO-MARKET STRATEGY', level=1)

phases = [
    ('Phase 1: Validation (Months 1-6)', ['Secure 3-5 pilot customers in progressive municipalities', 'Prove ROI with documented case studies', 'Refine product based on government user feedback']),
    ('Phase 2: Early Traction (Months 7-18)', ['Convert pilots to paid contracts', 'Expand to 20-30 facilities', 'Achieve government marketplace certifications', 'Build reference customer base']),
    ('Phase 3: Scale (Months 19-36)', ['Hire dedicated government sales team', 'Expand to state-level opportunities', 'Launch partner channel program', 'Pursue federal certifications (FedRAMP)'])
]

for phase_title, items in phases:
    doc.add_paragraph().add_run(phase_title).bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# Section 14: Risks and Mitigations
doc.add_heading('14. RISKS AND MITIGATIONS', level=1)

risk_table = doc.add_table(rows=6, cols=4)
risk_table.style = 'Table Grid'

risk_data = [
    ('Risk', 'Probability', 'Impact', 'Mitigation'),
    ('Long government sales cycles', 'High', 'Medium', 'Free pilots, strong ROI documentation'),
    ('Data security concerns', 'Medium', 'High', 'Government-grade security certifications'),
    ('Integration complexity', 'Medium', 'Medium', 'Pre-built connectors, professional services'),
    ('Budget cuts / political changes', 'Medium', 'High', 'Multi-year contracts, diversified customer base'),
    ('Competition from incumbents', 'Medium', 'Medium', 'Focus on AI differentiation and agility')
]

for i, (risk, prob, impact, mit) in enumerate(risk_data):
    row = risk_table.rows[i]
    row.cells[0].text = risk
    row.cells[1].text = prob
    row.cells[2].text = impact
    row.cells[3].text = mit
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Conclusion
doc.add_heading('CONCLUSION', level=1)

conclusion = doc.add_paragraph()
conclusion.add_run('GreenGrid AI addresses a critical intersection of public sector needs: reducing operational costs, meeting climate commitments, and modernizing aging infrastructure. By leveraging AI and data science to transform raw energy data into actionable intelligence, the platform delivers measurable value that justifies its cost within months, not years.')

doc.add_paragraph()

doc.add_paragraph('The combination of:')
combo = ['A clearly defined problem (energy waste in public buildings)', 'A defensible solution (AI-powered optimization)', 'A receptive market (government sustainability mandates)', 'A scalable business model (SaaS with network effects)']
for c in combo:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('...creates a compelling opportunity to build a category-defining company in the GovTech sustainability space.')

doc.add_paragraph()

# Market Validation Summary (Required by assignment - Max 500 words)
doc.add_page_break()
doc.add_heading('MARKET VALIDATION SUMMARY', level=1)

doc.add_paragraph().add_run('Research Methodology').bold = True
doc.add_paragraph('To validate the GreenGrid AI business concept, primary research was conducted through interviews with potential users and domain experts in the public sector energy management space. This research aimed to understand current pain points, existing solutions, and receptivity to AI-powered alternatives.')

doc.add_paragraph()
doc.add_paragraph().add_run('Interview Participants').bold = True

participants = [
    ('Participant A - Facilities Manager, Municipal Government', 'A facilities manager overseeing 12 public buildings including libraries, community centres, and administrative offices. 15+ years of experience in building operations and energy management.'),
    ('Participant B - Sustainability Officer, County Council', 'Responsible for implementing the county\'s Climate Action Plan and tracking progress toward 2030 net-zero commitments. Background in environmental policy and data analysis.'),
    ('Participant C - Energy Consultant, Public Sector Specialist', 'Independent consultant who has worked with over 20 local authorities on energy efficiency projects. Provides external perspective on market needs and competitive landscape.')
]

for title, desc in participants:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(desc, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph().add_run('Key Insights').bold = True

insights = [
    ('Pain Point Validation', 'All three participants confirmed that energy waste in public buildings is a significant and ongoing challenge. Participant A noted: "We know we\'re wasting energy, but we don\'t have the tools to identify where or when. By the time we get quarterly reports, the opportunity to act has passed." This validates our core problem hypothesis.'),
    ('Current Solutions Are Inadequate', 'Existing alternatives were described as either too expensive (external audits), too limited (basic BMS systems), or too manual (spreadsheet tracking). Participant B expressed frustration: "We have sustainability targets but no way to measure real-time progress. Everything is retrospective."'),
    ('AI Receptivity', 'There was strong interest in AI-powered solutions, particularly for predictive capabilities. Participant C noted that while governments are traditionally slow adopters, climate pressure is accelerating technology adoption: "Councils that declared climate emergencies are actively seeking innovative solutions."'),
    ('Key Concerns Identified', 'Data security and integration complexity emerged as primary concerns. Participant A emphasized: "Any solution must work with our existing systems—we cannot rip and replace." This insight directly informed our software-only, integration-focused approach.'),
    ('Willingness to Pay', 'All participants indicated budget availability for solutions with demonstrable ROI. Participant B confirmed: "If you can show 15-20% energy savings, the business case writes itself. That\'s money that goes back into public services."')
]

for title, insight in insights:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(insight)

doc.add_paragraph()
doc.add_paragraph().add_run('Conclusion').bold = True
doc.add_paragraph('The market validation research strongly supports the GreenGrid AI value proposition. The identified problems are real, current solutions are insufficient, and there is clear demand for an AI-powered platform that delivers measurable energy savings without requiring hardware investment. The insights gathered have directly shaped the product features, particularly the emphasis on integration capabilities, transparent AI recommendations, and ROI-focused reporting.')

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Document prepared for: CSC1112 Domain Applications 3 - Lean Canvas Assignment').italic = True

# Save
doc.save(r'D:\Final Year\App Domains 3\lean canvas idea\Lean_Business_Model_Canvas_Updated.docx')
print('Word document created successfully!')
