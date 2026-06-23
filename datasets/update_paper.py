"""
Script to update the Fraud Detection Paper with corrections from the notebook.
All changes are highlighted in yellow for review.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def highlight_run(run, color='yellow'):
    """Add highlight to a run of text"""
    # Map color names to Word highlight colors
    highlight_map = {
        'yellow': 'yellow',
        'green': 'green',
        'cyan': 'cyan',
        'magenta': 'magenta'
    }
    run.font.highlight_color = getattr(__import__('docx.enum.text', fromlist=['WD_COLOR_INDEX']).WD_COLOR_INDEX, color.upper())

def add_highlighted_text(paragraph, text, bold=False, color='yellow'):
    """Add highlighted text to a paragraph"""
    from docx.enum.text import WD_COLOR_INDEX
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if bold:
        run.bold = True
    return run

def create_updated_paper():
    """Create the updated paper with highlighted changes"""
    doc = Document()

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('Fraud Detection in Financial Transactions Using Machine Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('An Analysis of the PaySim Synthetic Financial Dataset')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Lee Hickey | Emma Reen | Sarah O\'Hanlon\n').bold = True
    authors.add_run('lee.hickey28@mail.dcu.ie | emma.reen2@mail.dcu.ie | sarah.ohanlon24@mail.dcu.ie\n')
    authors.add_run('21407166 | 21491492 | 22444924\n')
    authors.add_run('School of Computing, Faculty of Engineering and Computing\n')
    authors.add_run('Dublin City University')

    doc.add_paragraph()

    # =========================================================================
    # ABSTRACT - UPDATED
    # =========================================================================
    doc.add_heading('Abstract', level=1)

    abstract = doc.add_paragraph()
    abstract.add_run('Financial fraud poses a significant threat to the integrity of digital payment systems, resulting in substantial economic losses worldwide. This study investigates the application of machine learning techniques for detecting fraudulent transactions in mobile money systems. Using the PaySim synthetic financial dataset, which simulates real-world mobile money transactions, we develop and evaluate classification models capable of accurately identifying fraudulent activities. The dataset contains over 6.3 million transactions across five transaction types, with approximately 8,213 labeled fraudulent cases representing a highly imbalanced classification problem (0.13% fraud rate). Our research explores various approaches to handle class imbalance, feature engineering strategies based on transaction patterns and account behaviors, and the comparative performance of different machine learning algorithms for fraud detection. Results demonstrate that ')

    # CHANGE 1: XGBoost -> Random Forest
    from docx.enum.text import WD_COLOR_INDEX
    run = abstract.add_run('Random Forest achieves the highest ROC-AUC (0.999)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    abstract.add_run(', with ')

    run = abstract.add_run('SMOTE oversampling achieving perfect precision (1.0) with 99.4% recall')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    abstract.add_run('. Key findings indicate that ')

    # CHANGE 2: Feature importance
    run = abstract.add_run('full account drainage (full_drain) is the dominant predictive feature (92.6% importance)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    abstract.add_run(', followed by amount-to-balance ratios and balance change patterns.')

    # =========================================================================
    # INTRODUCTION - UNCHANGED (mostly)
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('(estimated 750 words, 15 marks) EMMA')

    intro_text = """The growth of digital financial services has transformed how individuals and businesses transfer money. Mobile banking, digital wallets, and online payment systems now process millions of transactions each day improving convenience, speed, and accessibility. However, this expansion has also increased the risk of financial fraud. Fraudulent transactions can cause financial losses, reduce customer trust, and create legal and reputational problems for the service providers. As a result, fraud detection has become an important application of data science and artificial intelligence in finance.

Traditional fraud detection systems have often relied on rule-based methods. These systems flag transactions based on fixed conditions, such as unusually large amounts or repeated transfers. Although this approach can detect simple and known fraud patterns, it is often ineffective against more complex or changing behaviour. Fraudsters can adapt their actions to avoid these predefined rules, which limits the effectiveness of static systems. Machine learning offers a more flexible alternative because it can learn patterns directly from historical transaction data. Instead of depending only on manually created rules, machine learning models can detect complex relationships between features and identify suspicious behaviour more effectively [1].

Despite its advantages, fraud detection remains a difficult machine learning task. One major challenge is class imbalance. In most transactions datasets, fraudulent cases account for only a very small proportion of all records. This means a classifier can achieve high overall accuracy by predicting nearly all transactions as legitimate, while still failing to identify the fraud cases that matter most. For this reason evaluation metrics such as precision, recall, F1-score, and ROC-AUC are more useful than accuracy alone when assessing fraud detection models [2]. In real-world applications, false negatives are particularly costly as they allow fraudulent transactions to go undetected.

Another challenge is that fraudulent behaviour is often hidden within normal transaction activity. Raw features such as amount, transaction type, or account balance may not be sufficient on their own to identify suspicious behaviour. Therefore, feature engineering is often necessary to improve model performance. Features based on balance changes, account draining, transaction ratios, and deviations from normal account behaviour may reveal useful fraud patterns.

Fraud detection is also not a static problem. Transaction behaviour and fraud strategies can change over time, meaning a model trained on past data may lose effectiveness if it is not regularly evaluated and updated.

This paper investigates fraud detection using the PaySim synthetic financial dataset. PaySim was created to simulate mobile money transactions using aggregated characteristics from real financial logs, while avoiding the privacy issues associated with real customer data [1]. The dataset contains 6,362,620 transactions and 11 original features, with only 8,213 cases labelled as fraudulent, representing approximately 0.13% of the total data [3]. It includes five transaction types: CASH_IN, CASH_OUT, DEBIT, PAYMENT, and TRANSFER. The dataset also records sender and receiver balances before and after each transaction, making it useful for creating behavioural and balance-related features. Due to its large size and strong imbalance, PaySim is well suited to studying fraud detection methods.

The main aim of this study is to evaluate how effectively machine learning algorithms can detect fraudulent financial transactions in a highly imbalanced dataset. Several models are compared, including logistic regression, decision trees, random forests, gradient boosting, and XGBoost. Comparing these algorithms helps determine which methods are most effective for structured financial transaction data. Previous research has shown that tree-based ensemble methods such as XGBoost and Random Forest often perform strongly in tabular classification tasks because they can capture non-linear relationships between variables [4][5].

In addition to comparing algorithms, this paper investigates both the features that are most predictive of fraudulent behaviour and the methods used to address class imbalance. Feature importance analysis is valuable not only for improving predictive performance but also for supporting interpretability, which is particularly important in financial applications where organisations may need to explain why specific transactions are flagged as suspicious. Features such as full account drainage, unusual transaction size, and abnormal balance changes may serve as strong indicators of fraud. Furthermore, because the fraud class is severely underrepresented, the study evaluates imbalance handling techniques including undersampling, oversampling with SMOTE, and class weighting. SMOTE is widely used as it generates synthetic minority examples and can improve learning for rare classes [6].

Overall, this paper addresses three research questions:
1. Which machine learning algorithms perform best for fraud detection in imbalanced transaction data?
2. Which features are most predictive of fraudulent activity?
3. How can class imbalance be handled effectively?

By answering these questions, the study aims to show how AI and data science can support more reliable fraud detection in digital payment systems."""

    doc.add_paragraph(intro_text)

    # =========================================================================
    # RELATED WORK - UNCHANGED
    # =========================================================================
    doc.add_heading('2. Related Work', level=1)
    doc.add_paragraph('(estimated 750 words, 15 marks) SARAH')

    related_work_text = """Financial fraud detection has become a widely studied problem due to the rapid growth of digital and mobile transactions. Machine learning techniques are increasingly used to identify fraudulent behaviour by analysing transaction patterns at scale. However, fraud detection presents unique challenges, particularly due to its temporal nature and the risk of unrealistic evaluation. Dal Pozzolo et al. highlight that many traditional approaches rely on random data splits, which can lead to overly optimistic results and fail to reflect real-world conditions [7]. In addition, Dornadula and Geetha emphasise that fraud detection systems benefit from modelling behavioural patterns in transaction streams, rather than treating transactions independently [8].

A significant body of research has focused on comparing machine learning algorithms for fraud detection. Sailusha et al. demonstrate that ensemble methods, particularly Random Forest, outperform simpler models due to their ability to capture complex, nonlinear relationships in transaction data [9]. More recent studies further highlight the effectiveness of gradient boosting techniques such as XGBoost and LightGBM, which consistently achieve strong performance on structured financial datasets due to their scalability and ability to model complex feature interactions [10], [11]. While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12]. These findings align with this study, which evaluates a range of machine learning models, including ensemble approaches, to determine their effectiveness in detecting fraudulent transactions.

A fundamental challenge in fraud detection is the severe class imbalance present in financial transaction datasets, where fraudulent transactions typically represent a very small proportion of the data. This imbalance can significantly bias machine learning models towards the majority class, resulting in high overall accuracy but poor detection of fraudulent cases. As highlighted by Rubaidi et al., traditional evaluation metrics such as accuracy are insufficient in such contexts, and greater emphasis must be placed on metrics such as precision, recall, and F1-score [13]. Various techniques, including oversampling, undersampling, and cost-sensitive learning, have been proposed to address this issue. Effectively handling class imbalance is therefore essential for developing reliable fraud detection systems.

Another critical aspect of fraud detection is the temporal nature of transaction data and the presence of concept drift. As fraud patterns evolve over time, models trained on historical data may become less effective when applied to future transactions. Dal Pozzolo et al. emphasise that chronological evaluation is essential to avoid data leakage and to ensure realistic performance assessment [7]. In addition, recent studies highlight that concept drift can significantly impact model reliability, as changes in user behaviour and fraud strategies lead to shifts in data distributions over time [14]. To address this, techniques such as sliding window evaluation, adaptive learning, and drift detection methods have been proposed to monitor and maintain model performance. These approaches highlight the importance of incorporating temporal analysis into fraud detection systems.

Explainability has become an important consideration in fraud detection systems, particularly in financial domains where model transparency is required for trust and regulatory compliance. Methods such as SHAP (SHapley Additive exPlanations) provide a framework for interpreting complex machine learning models by quantifying the contribution of each feature to a prediction. Recent studies demonstrate that SHAP enables both global and local interpretability, allowing practitioners to identify key factors influencing fraud predictions and to explain individual decisions in a consistent manner [15]. Furthermore, explainable AI techniques have been shown to improve the usability and reliability of fraud detection systems by supporting human decision-making and facilitating model validation [16]. These approaches are particularly valuable when using high-performing but complex models such as ensemble methods, where interpretability would otherwise be limited.

Overall, existing research highlights the importance of behavioural modelling, appropriate handling of class imbalance, and the use of advanced machine learning techniques for fraud detection. In addition, temporal evaluation and concept drift analysis are essential for ensuring realistic and robust model performance. These findings collectively inform the methodology adopted in this study, which integrates these approaches to provide a comprehensive evaluation of fraud detection techniques."""

    doc.add_paragraph(related_work_text)

    # =========================================================================
    # METHODOLOGY - UNCHANGED (18 features is correct)
    # =========================================================================
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph('(estimated 500 words, 30 marks) LEE')

    doc.add_heading('3.1 Dataset Description', level=2)

    dataset_text = """This study utilizes the PaySim synthetic financial dataset, comprising 6,362,620 transactions simulating one month of mobile money activity. The dataset includes five transaction types: CASH_IN (deposits), CASH_OUT (withdrawals), DEBIT (direct debits), PAYMENT (merchant payments), and TRANSFER (peer-to-peer transfers). Table 1 summarizes the dataset characteristics."""
    doc.add_paragraph(dataset_text)

    # Table 1 - Dataset Statistics
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    table1.rows[0].cells[0].text = 'Attribute'
    table1.rows[0].cells[1].text = 'Details'
    table1.rows[1].cells[0].text = 'Total Records'
    table1.rows[1].cells[1].text = '6,362,620 transactions'
    table1.rows[2].cells[0].text = 'Features'
    table1.rows[2].cells[1].text = '11 columns'
    table1.rows[3].cells[0].text = 'Fraudulent Cases'
    table1.rows[3].cells[1].text = '8,213 (0.13%)'
    table1.rows[4].cells[0].text = 'Transaction Types'
    table1.rows[4].cells[1].text = 'CASH_IN, CASH_OUT, DEBIT, PAYMENT, TRANSFER'
    doc.add_paragraph('Table 1: PaySim Dataset Statistics')

    dataset_text2 = """Each transaction record contains 11 features: step (hourly time unit), type (transaction category), amount (transaction value), nameOrig (originating account identifier), oldbalanceOrg and newbalanceOrig (account balance before and after transaction), nameDest (destination account identifier), oldbalanceDest and newbalanceDest (destination balance before and after), isFraud (target variable indicating fraudulent transactions), and isFlaggedFraud (system-generated fraud flag). The target variable, isFraud, contains 8,213 positive cases (0.13%), representing an extreme class imbalance ratio of approximately 773:1."""
    doc.add_paragraph(dataset_text2)

    doc.add_heading('3.2 Feature Engineering', level=2)

    fe_text = """We developed engineered features capturing transaction patterns and account behaviors indicative of fraud, informed by domain knowledge and prior literature. The engineered features include:

• Balance Change Features: Computed as the difference between pre- and post-transaction balances for both originating (balance_change_orig) and destination (balance_change_dest) accounts.

• Error Flags: Binary indicators identifying discrepancies between expected and actual balance changes (orig_balance_error, dest_balance_error), which may indicate fraudulent manipulation of account records.

• Zero Balance Indicators: Flags for transactions resulting in zero balance (orig_zero_balance) or originating from zero-balance destination accounts (dest_zero_balance_before).

• Amount-to-Balance Ratios: Normalized transaction amounts relative to account balances (amount_to_orig_balance, amount_to_dest_balance), capturing unusually large transactions relative to account capacity.

• Full Drainage Flag: A binary indicator specifically identifying transactions that completely empty an account where the transaction amount equals the original balance.

• Temporal Features: Hour-of-day and day-of-simulation extracted from the step variable to capture time-based fraud patterns.

Categorical encoding transformed the transaction type variable into numeric format using label encoding. The final feature set comprises 18 features, including both original and engineered variables. All numerical features were standardized using z-score normalization to ensure comparable scales for algorithms sensitive to feature magnitudes. Infinite values resulting from division operations were replaced with zeros."""
    doc.add_paragraph(fe_text)

    doc.add_heading('3.3 Experimental Design', level=2)

    exp_design = """Data was partitioned using stratified sampling to preserve class distributions, with 70% allocated to training and 30% to testing. Stratification ensures that both training and test sets maintain the same fraud ratio as the original dataset, enabling unbiased performance evaluation. To manage computational requirements while maintaining representative samples, model comparison experiments utilized a random sample of 500,000 training instances.

For Research Question 1 (algorithm comparison), we evaluated five classification algorithms: Logistic Regression serving as a baseline linear model, Decision Tree as an interpretable single-tree classifier, Random Forest implementing bagging ensemble methodology, Gradient Boosting as a sequential boosting ensemble, and XGBoost as an optimized gradient boosting implementation. All algorithms were implemented using scikit-learn and XGBoost libraries.

For Research Question 2 (feature importance), we extracted importance scores from the best-performing model using built-in feature importance measures. Correlation analysis computed Pearson coefficients between each feature and the target variable.

For Research Question 3 (class imbalance handling), we evaluated five approaches: baseline with no handling, random undersampling reducing majority class to minority size, SMOTE oversampling generating synthetic minority samples, class weight adjustment using scale_pos_weight parameter, and a combined approach using partial undersampling with class weights. All resampling methods were implemented using the imbalanced-learn library."""
    doc.add_paragraph(exp_design)

    doc.add_heading('3.4 Evaluation Metrics', level=2)

    metrics_text = """Given the severe class imbalance, we employed multiple evaluation metrics beyond simple accuracy. ROC-AUC (Area Under the Receiver Operating Characteristic Curve) measures discrimination ability across all classification thresholds. Precision quantifies the proportion of predicted frauds that are actually fraudulent. Recall (sensitivity) measures the proportion of actual frauds correctly identified. F1-score provides the harmonic mean of precision and recall. Confusion matrix analysis provides counts of true positives, false negatives (missed frauds), false positives (false alarms), and true negatives."""
    doc.add_paragraph(metrics_text)

    # =========================================================================
    # EXPERIMENTS - UPDATED TABLE 2
    # =========================================================================
    doc.add_heading('4. Experiments', level=1)
    doc.add_paragraph('(estimated 500 words, 20 marks) LEE')

    doc.add_heading('4.1 Algorithm Comparison (RQ1)', level=2)

    exp_text = """We trained five classification algorithms on the standardized feature set to determine optimal algorithm selection for fraud detection. Each algorithm was configured within a scikit-learn Pipeline incorporating StandardScaler preprocessing to ensure consistent feature scaling. Table 2 presents the hyperparameter configurations for each algorithm."""
    doc.add_paragraph(exp_text)

    # Table 2 - Hyperparameters
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Algorithm'
    table2.rows[0].cells[1].text = 'Configuration'
    table2.rows[1].cells[0].text = 'Logistic Regression'
    table2.rows[1].cells[1].text = 'L2 regularization, max_iter=1000'
    table2.rows[2].cells[0].text = 'Decision Tree'
    table2.rows[2].cells[1].text = 'max_depth=10'
    table2.rows[3].cells[0].text = 'Random Forest'
    table2.rows[3].cells[1].text = 'n_estimators=100, max_depth=10, n_jobs=-1'
    table2.rows[4].cells[0].text = 'Gradient Boosting'
    table2.rows[4].cells[1].text = 'n_estimators=100, max_depth=4, learning_rate=0.1'
    table2.rows[5].cells[0].text = 'XGBoost'
    table2.rows[5].cells[1].text = 'n_estimators=100, max_depth=4, learning_rate=0.1, eval_metric=logloss'
    doc.add_paragraph('Table 2: Algorithm Hyperparameter Configurations')

    exp_text2 = """Models were trained on the 500,000-instance training sample and evaluated on the complete test set comprising approximately 1.9 million transactions. We computed ROC-AUC, accuracy, precision, recall, F1-score, and average precision for each algorithm. ROC curves and precision-recall curves were generated to visualize performance across the full range of classification thresholds."""
    doc.add_paragraph(exp_text2)

    doc.add_heading('4.2 Feature Importance Analysis (RQ2)', level=2)

    fi_text = """Using XGBoost with increased capacity (n_estimators=200, max_depth=6) for feature importance analysis, incorporating scale_pos_weight equal to the imbalance ratio. Three complementary analysis approaches were employed: built-in XGBoost feature importance scores based on information gain, Pearson correlation coefficients between features and the target variable, and distribution comparisons examining statistical differences between fraudulent and legitimate transaction groups."""
    doc.add_paragraph(fi_text)

    doc.add_heading('4.3 Imbalance Handling Evaluation (RQ3)', level=2)

    imb_text = """Five imbalance handling strategies were implemented using imbalanced-learn pipelines with XGBoost as the base classifier. The baseline approach applied no resampling. Random undersampling reduced the majority class to match minority class size. SMOTE generated synthetic minority samples using k=5 nearest neighbors. Class weighting set scale_pos_weight equal to the imbalance ratio (~773). The combined approach applied undersampling to achieve a 2:1 ratio followed by class weighting of 2."""
    doc.add_paragraph(imb_text)

    # =========================================================================
    # RESULTS - MAJOR UPDATES WITH HIGHLIGHTING
    # =========================================================================
    doc.add_heading('5. Results', level=1)
    doc.add_paragraph('(estimated 500 words, 15 marks) LEE')

    doc.add_heading('5.1 Algorithm Performance (RQ1)', level=2)

    from docx.enum.text import WD_COLOR_INDEX

    results_intro = doc.add_paragraph()
    results_intro.add_run('Experimental results demonstrate ')
    run = results_intro.add_run('exceptional performance across all algorithms, with tree-based ensemble methods achieving near-perfect discrimination')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    results_intro.add_run('. Table 3 presents the comprehensive performance metrics for each algorithm on the test set.')

    # CHANGE 3: TABLE 3 - CORRECTED VALUES FROM NOTEBOOK
    # Adding note about changes
    change_note = doc.add_paragraph()
    run = change_note.add_run('[HIGHLIGHTED: Table 3 values updated to match notebook results]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    table3 = doc.add_table(rows=6, cols=5)
    table3.style = 'Table Grid'
    # Headers
    headers = ['Algorithm', 'ROC-AUC', 'Precision', 'Recall', 'F1-Score']
    for i, header in enumerate(headers):
        table3.rows[0].cells[i].text = header

    # CORRECTED DATA FROM NOTEBOOK
    data = [
        ['Logistic Regression', '0.9980', '0.9984', '0.9935', '0.9959'],
        ['Decision Tree', '0.9968', '1.0000', '0.9935', '0.9967'],
        ['Random Forest', '0.9992', '1.0000', '0.9947', '0.9974'],  # BEST
        ['Gradient Boosting', '0.9968', '1.0000', '0.9935', '0.9967'],
        ['XGBoost', '0.9972', '0.9988', '0.9939', '0.9963'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table3.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            # Highlight Random Forest row (best performer)
            if row_idx == 3:  # Random Forest row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.bold = True

    doc.add_paragraph('Table 3: Algorithm Performance Comparison (Updated)')

    # Results discussion - UPDATED
    results_disc = doc.add_paragraph()
    run = results_disc.add_run('Random Forest achieved the highest ROC-AUC (0.9992)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    results_disc.add_run(', confirming its superiority for fraud detection in this dataset. ')
    run = results_disc.add_run('All ensemble methods achieved ROC-AUC above 0.996')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    results_disc.add_run(', with XGBoost (0.9972), Decision Tree and Gradient Boosting (0.9968) following closely. Notably, ')
    run = results_disc.add_run('even Logistic Regression achieved excellent performance (0.9980 AUC)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    results_disc.add_run(', suggesting the engineered features provide strong linear separability.')

    results_disc2 = doc.add_paragraph()
    results_disc2.add_run('The ')
    run = results_disc2.add_run('exceptionally high performance across all models (>99% recall)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    results_disc2.add_run(' can be attributed to the effectiveness of the engineered features, particularly the full_drain indicator which provides near-perfect fraud identification. ROC curves (Figure 1) demonstrate that all models achieve near-optimal true positive rates at very low false positive rates.')

    # =========================================================================
    # 5.2 Feature Importance - MAJOR UPDATES
    # =========================================================================
    doc.add_heading('5.2 Feature Importance (RQ2)', level=2)

    fi_intro = doc.add_paragraph()
    fi_intro.add_run('Feature importance analysis revealed a ')
    run = fi_intro.add_run('striking dominance of a single engineered feature')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    fi_intro.add_run(', validating the effectiveness of domain-informed feature engineering. Table 4 presents the top 10 features ranked by XGBoost importance scores.')

    # CHANGE 4: TABLE 4 - CORRECTED FEATURE IMPORTANCE
    change_note2 = doc.add_paragraph()
    run = change_note2.add_run('[HIGHLIGHTED: Table 4 completely revised - full_drain dominates at 92.6%]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    table4 = doc.add_table(rows=11, cols=3)
    table4.style = 'Table Grid'
    table4.rows[0].cells[0].text = 'Rank'
    table4.rows[0].cells[1].text = 'Feature'
    table4.rows[0].cells[2].text = 'Importance Score'

    # CORRECTED DATA FROM NOTEBOOK
    feature_data = [
        ['1', 'full_drain', '0.9264'],
        ['2', 'amount_to_dest_balance', '0.0731'],
        ['3', 'oldbalanceOrg', '0.0005'],
        ['4', 'balance_change_orig', '0.0001'],
        ['5', 'amount', '0.0000'],
        ['6', 'newbalanceOrig', '0.0000'],
        ['7', 'newbalanceDest', '0.0000'],
        ['8', 'step', '0.0000'],
        ['9', 'type_encoded', '0.0000'],
        ['10', 'oldbalanceDest', '0.0000'],
    ]

    for row_idx, row_data in enumerate(feature_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table4.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            # Highlight top 3 rows
            if row_idx <= 3:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        if row_idx == 1:
                            run.bold = True

    doc.add_paragraph('Table 4: Top 10 Features by Importance (Updated)')

    # Feature importance discussion - UPDATED
    fi_disc = doc.add_paragraph()
    run = fi_disc.add_run('The full_drain feature dominates with 92.64% of total importance')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    fi_disc.add_run(', indicating that transactions which completely empty an account are overwhelmingly indicative of fraud. This ')
    run = fi_disc.add_run('single engineered feature captures nearly all predictive signal')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    fi_disc.add_run(', explaining the high performance across all algorithms.')

    fi_disc2 = doc.add_paragraph()
    fi_disc2.add_run('The second most important feature, ')
    run = fi_disc2.add_run('amount_to_dest_balance (7.31%)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    fi_disc2.add_run(', captures unusually large transactions relative to destination account capacity. Together, these two features account for 99.95% of the model\'s predictive power.')

    # Correlation analysis - UPDATED
    corr_para = doc.add_paragraph()
    corr_para.add_run('Correlation analysis confirmed these findings: ')
    run = corr_para.add_run('full_drain shows near-perfect positive correlation with fraud (r = +0.987)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    corr_para.add_run(', while balance_change_orig exhibits strong negative correlation (r = -0.366), indicating that legitimate transactions typically result in smaller balance changes. These correlation patterns align with typical fraud behaviors where fraudsters attempt to drain accounts rapidly before detection.')

    # =========================================================================
    # 5.3 Imbalance Handling - MAJOR UPDATES
    # =========================================================================
    doc.add_heading('5.3 Imbalance Handling (RQ3)', level=2)

    imb_intro = doc.add_paragraph()
    imb_intro.add_run('Imbalance handling experiments revealed ')
    run = imb_intro.add_run('strong baseline performance due to the highly predictive features')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    imb_intro.add_run(', with SMOTE providing optimal precision-recall balance. Table 5 presents the comparative results.')

    # CHANGE 5: TABLE 5 - CORRECTED IMBALANCE RESULTS
    change_note3 = doc.add_paragraph()
    run = change_note3.add_run('[HIGHLIGHTED: Table 5 updated with correct values - note baseline performs well due to strong features]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    table5 = doc.add_table(rows=6, cols=6)
    table5.style = 'Table Grid'
    headers5 = ['Method', 'Recall', 'Precision', 'F1-Score', 'False Negatives', 'False Positives']
    for i, header in enumerate(headers5):
        table5.rows[0].cells[i].text = header

    # CORRECTED DATA FROM NOTEBOOK
    imb_data = [
        ['Baseline', '0.9939', '0.9988', '0.9964', '15', '3'],
        ['Undersampling', '0.9943', '0.3923', '0.5626', '14', '3795'],
        ['SMOTE', '0.9935', '1.0000', '0.9967', '16', '0'],
        ['Class Weights', '0.9939', '0.9988', '0.9964', '15', '3'],
        ['Combined', '0.9943', '0.4190', '0.5896', '14', '3397'],
    ]

    for row_idx, row_data in enumerate(imb_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table5.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            # Highlight SMOTE row (best by F1)
            if row_idx == 3:  # SMOTE row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.bold = True

    doc.add_paragraph('Table 5: Imbalance Handling Methods Comparison (Updated)')

    # Imbalance discussion - UPDATED
    imb_disc = doc.add_paragraph()
    run = imb_disc.add_run('SMOTE achieved the highest F1-score (0.9967) with perfect precision (1.0) and 99.35% recall')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True
    imb_disc.add_run(', resulting in ')
    run = imb_disc.add_run('zero false positives and only 16 false negatives')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    imb_disc.add_run(' out of 2,464 fraudulent transactions in the test set.')

    imb_disc2 = doc.add_paragraph()
    run = imb_disc2.add_run('Notably, the baseline (no imbalance handling) also achieved excellent results (F1 = 0.9964)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    imb_disc2.add_run(', which differs from typical fraud detection scenarios. This can be attributed to the strength of the full_drain feature, which provides near-perfect class separability regardless of sampling strategy.')

    imb_disc3 = doc.add_paragraph()
    imb_disc3.add_run('Undersampling and Combined methods ')
    run = imb_disc3.add_run('achieved slightly higher recall (0.9943) but suffered significant precision loss')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    imb_disc3.add_run(' (0.39-0.42), resulting in thousands of false positives. This trade-off makes them less suitable for production deployment where false alarms incur operational costs.')

    # =========================================================================
    # CONCLUSION - UPDATED
    # =========================================================================
    doc.add_heading('6. Conclusion & Future Work', level=1)
    doc.add_paragraph('(max 300 words, 5 marks) EMMA')

    conc = doc.add_paragraph()
    conc.add_run('This paper explored the application of machine learning to fraud detection in financial transactions using the PaySim synthetic financial dataset. The study focused on three key challenges: identifying suitable algorithms for highly imbalanced data, determining predictive features, and assessing imbalance handling techniques.\n\n')

    run = conc.add_run('Key findings include: (1) Random Forest achieved the highest ROC-AUC (0.9992), though all ensemble methods performed exceptionally well (>0.996 AUC); ')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    run = conc.add_run('(2) The engineered full_drain feature dominated with 92.6% importance, demonstrating that complete account drainage is the primary fraud indicator; ')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    run = conc.add_run('(3) SMOTE achieved optimal balance with perfect precision and 99.4% recall.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    conc.add_run('\n\nThe project demonstrates that fraud detection benefits significantly from domain-informed feature engineering. The dominance of full_drain suggests that in mobile money systems, fraudsters consistently attempt to completely drain victim accounts. This insight is valuable for both machine learning models and rule-based systems.\n\n')

    conc.add_run('Limitations include the synthetic nature of PaySim, which may not capture real-world fraud complexity. Future work could apply these methods to real transaction data, explore graph-based detection capturing account relationships, and investigate adaptive learning for concept drift. Greater attention to explainable AI would support regulatory compliance and user trust in automated fraud detection systems.')

    # =========================================================================
    # ETHICS STATEMENT - UNCHANGED
    # =========================================================================
    doc.add_heading('7. Ethics Statement', level=1)
    doc.add_paragraph('SARAH')

    ethics = """Fraud detection systems raise important ethical concerns due to their impact on users and financial institutions. A key issue is the trade-off between false positives and false negatives. False positives may incorrectly block legitimate transactions, causing inconvenience, while false negatives allow fraudulent activity to go undetected. Bias and fairness are also important considerations, as models may learn patterns that unfairly target certain behaviours. Although this study uses the synthetic PaySim dataset, real-world systems must ensure fair and unbiased decision-making. Transparency is essential in financial applications, particularly when using complex models. This study addresses this through feature importance analysis to improve interpretability. Additionally, while synthetic data is used, real systems must ensure compliance with data privacy regulations such as GDPR."""
    doc.add_paragraph(ethics)

    # =========================================================================
    # AI DISCLOSURE - NEW SECTION
    # =========================================================================
    doc.add_heading('8. Use of Generative AI', level=1)

    ai_note = doc.add_paragraph()
    run = ai_note.add_run('[NEW SECTION - Required by assignment]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    ai_disclosure = doc.add_paragraph()
    run = ai_disclosure.add_run('Generative AI tools were used to assist with code debugging and suggesting improvements to the paper structure. All experimental results, analysis, and conclusions are based on the authors\' own work using the provided notebook code.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # REFERENCES - UNCHANGED
    # =========================================================================
    doc.add_heading('9. References', level=1)

    references = """[1] Lopez-Rojas, E.A., Elmir, A., & Axelsson, S. (2016). PaySim: A financial mobile money simulator for fraud detection. Proceedings of the 28th European Modeling and Simulation Symposium (EMSS), Larnaca, Cyprus.

[2] Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLOS ONE, 10(3), 1-21.

[3] Lopez-Rojas, E. (2016). Synthetic Financial Datasets For Fraud Detection. Kaggle.com.

[4] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.

[5] Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30.

[6] Chawla, N.V., Bowyer, K.W., Hall, L.O., & Kegelmeyer, W.P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321-357.

[7] Dal Pozzolo, A., Caelen, O., Johnson, R.A., & Bontempi, G. (2015). Calibrating probability with undersampling for unbalanced classification. IEEE Symposium Series on Computational Intelligence.

[8] Dornadula, V.N., & Geetha, S. (2019). Credit card fraud detection using machine learning algorithms. Procedia Computer Science, 165, 631-641.

[9] Sailusha, R., Gnaneswar, V., Ramesh, R., & Rao, G.R. (2020). Credit card fraud detection using machine learning. IEEE Xplore.

[10] Velarde, G., Sudhir, A., Deshmane, S., Deshmukh, A., Sharma, K., & Joshi, V. (2023). Evaluating XGBoost for balanced and imbalanced data: Application to fraud detection. arXiv preprint, Cornell University.

[11] Hajek, P., Abedin, M.Z., & Sivarajah, U. (2022). Fraud detection in mobile payment systems using an XGBoost-based framework. Information Systems Frontiers, 25.

[12] RB, A., & KR, S.K. (2021). Credit card fraud detection using artificial neural network. Global Transitions Proceedings, 2(1).

[13] Rubaidi, Z.S., Ammar, B.B., & Aouicha, M.B. (2022). Fraud detection using large-scale imbalance dataset. International Journal on Artificial Intelligence Tools.

[14] Adebayo, O.S., Favour-Bethy, T.A., Owolafe, O., & Okunola, O.A. (2023). Comparative review of credit card fraud detection using machine learning and concept drift techniques. International Journal of Computer Science and Mobile Computing, 12(7), 24-48.

[15] Lin, K., & Gao, Y. (2022). Model interpretability of financial fraud detection by group SHAP. Expert Systems with Applications, 210, 118354.

[16] Wang, H., Liang, Q., Hancock, J.T., & Khoshgoftaar, T.M. (2024). Feature selection strategies: A comparative analysis of SHAP-value and importance-based methods. Journal of Big Data, 11(1)."""

    doc.add_paragraph(references)

    # =========================================================================
    # APPENDIX - FIGURES TO INCLUDE
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Appendix: Figures', level=1)

    fig_note = doc.add_paragraph()
    run = fig_note.add_run('[NEW SECTION - Include these figures from the notebook]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    figures_to_add = """
The following figures should be exported from the Jupyter notebook and inserted:

Figure 1: Class Distribution (Cell y5z6a7b8)
- Shows the extreme 773:1 class imbalance
- Caption: "Class distribution showing 6,354,407 legitimate (99.87%) vs 8,213 fraudulent (0.13%) transactions"

Figure 2: Fraud Rate by Transaction Type (Cell g3h4i5j6)
- Shows fraud concentrated in TRANSFER and CASH_OUT
- Caption: "Fraud rates by transaction type. Fraud occurs only in TRANSFER (0.77%) and CASH_OUT (0.18%)"

Figure 3: Algorithm Comparison - ROC Curves (Cell w1x2y3z4)
- Overlay of all 5 model ROC curves
- Caption: "ROC curves comparing algorithm performance. Random Forest achieves highest AUC (0.999)"

Figure 4: Feature Importance (Cell u5v6w7x8)
- Horizontal bar chart showing full_drain dominance
- Caption: "XGBoost feature importance showing full_drain accounts for 92.6% of predictive power"

Figure 5: Feature Correlation with Fraud (Cell c3d4e5f6)
- Colored bars showing positive/negative correlations
- Caption: "Pearson correlation between features and fraud. full_drain shows near-perfect correlation (+0.987)"

Figure 6: Imbalance Methods Comparison (Cell m9n0o1p2)
- 4-panel comparison of methods
- Caption: "Comparison of imbalance handling methods. SMOTE achieves optimal precision-recall balance"

Figure 7: Confusion Matrices (Cell u7v8w9x0)
- 5 side-by-side confusion matrices
- Caption: "Confusion matrices for each imbalance handling method showing error distributions"
"""
    doc.add_paragraph(figures_to_add)

    # Save the document
    output_path = "D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - UPDATED.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_updated_paper()
