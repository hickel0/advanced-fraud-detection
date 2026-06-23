"""
Complete fix for the Fraud Detection Paper to align with notebook results.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# Load the document
doc = Document('D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper (1).docx')

changes_log = []

def find_and_replace_text(doc, old_text, new_text, description):
    """Find and replace text across all paragraphs."""
    found = False
    for para in doc.paragraphs:
        if old_text in para.text:
            # Get full paragraph text
            full_text = para.text
            new_full_text = full_text.replace(old_text, new_text)

            # Clear all runs and set new text
            if para.runs:
                # Preserve first run's formatting
                first_run_font = para.runs[0].font
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_full_text
            found = True
            changes_log.append(f"FIXED: {description}")

    return found

# ============================================================================
# FIX 1: Introduction - Remove LightGBM and neural networks from algorithm list
# ============================================================================
find_and_replace_text(
    doc,
    "Several models are compared, including logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks.",
    "Several models are compared, including Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost.",
    "Introduction: Removed LightGBM and neural networks from algorithm list"
)

# Also try with different capitalization
find_and_replace_text(
    doc,
    "logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks",
    "Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost",
    "Introduction: Fixed algorithm list capitalization"
)

# ============================================================================
# FIX 2: Related Work - Remove neural network sentence
# ============================================================================
find_and_replace_text(
    doc,
    "While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12].",
    "",
    "Related Work: Removed neural network comparison sentence"
)

# ============================================================================
# FIX 3: Related Work - Fix "These findings align" sentence
# ============================================================================
find_and_replace_text(
    doc,
    "including ensemble approaches and neural networks, to determine their effectiveness",
    "including tree-based ensemble methods, to determine their effectiveness",
    "Related Work: Changed 'ensemble approaches and neural networks' to 'tree-based ensemble methods'"
)

# ============================================================================
# FIX 4: Related Work - Fix LightGBM mention in "More recent studies" sentence
# ============================================================================
find_and_replace_text(
    doc,
    "More recent studies further highlight the effectiveness of gradient boosting techniques such as XGBoost and LightGBM, which consistently achieve strong performance on structured financial datasets due to their scalability and ability to model complex feature interactions [10], [11].",
    "More recent studies further highlight the effectiveness of gradient boosting techniques, particularly XGBoost, which consistently achieves strong performance on structured financial datasets due to its scalability and ability to model complex feature interactions [10].",
    "Related Work: Removed LightGBM from gradient boosting techniques discussion"
)

# ============================================================================
# FIX 5: Introduction - Fix the "primary aim" paragraph with LightGBM mention
# ============================================================================
# First, let's see the exact text
for para in doc.paragraphs:
    if "LightGBM" in para.text and "primary aim" in para.text.lower():
        print(f"Found paragraph with LightGBM: {para.text[:200]}...")

# Replace in primary aim paragraph
find_and_replace_text(
    doc,
    "Several models are compared, including logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks. Comparing these algorithms helps determine which methods are most effective for structured financial transaction data.",
    "Several models are compared, including Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost. Comparing these algorithms helps determine which methods are most effective for structured financial transaction data.",
    "Introduction: Fixed algorithm comparison sentence"
)

# ============================================================================
# Now scan for any remaining issues
# ============================================================================
print("\n" + "="*70)
print("SCANNING FOR REMAINING ISSUES...")
print("="*70)

remaining_issues = []
para_idx = 0

for para in doc.paragraphs:
    para_idx += 1
    text_lower = para.text.lower()

    # Skip empty paragraphs
    if not para.text.strip():
        continue

    # Check for LightGBM (skip references section)
    if 'lightgbm' in text_lower:
        # Check if this is in references (by checking for typical reference patterns)
        if not ('ke, g.' in text_lower or '2017' in para.text):
            remaining_issues.append({
                'para': para_idx,
                'issue': 'LightGBM reference',
                'text': para.text[:150] + '...' if len(para.text) > 150 else para.text
            })

    # Check for neural network (skip references and future work)
    if 'neural network' in text_lower:
        # Skip if in references (contains author names pattern)
        if 'rb, a.' in text_lower or 'artificial neural network' in text_lower:
            continue
        # Skip if in future work section (graph neural networks is OK)
        if 'graph neural' in text_lower or 'future' in text_lower:
            continue
        remaining_issues.append({
            'para': para_idx,
            'issue': 'Neural network reference',
            'text': para.text[:150] + '...' if len(para.text) > 150 else para.text
        })

# ============================================================================
# Additional targeted fixes based on remaining issues
# ============================================================================

# Fix any remaining algorithm lists
for para in doc.paragraphs:
    if 'LightGBM' in para.text and 'XGBoost' in para.text and 'neural' in para.text.lower():
        # This is likely the algorithm list, fix it directly
        new_text = para.text.replace('XGBoost, LightGBM, and neural networks', 'and XGBoost')
        new_text = new_text.replace('XGBoost, LightGBM and neural networks', 'and XGBoost')
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        changes_log.append("Fixed remaining algorithm list with LightGBM and neural networks")

# ============================================================================
# Save the document
# ============================================================================
output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - CORRECTED.docx'
doc.save(output_path)

# ============================================================================
# Print summary
# ============================================================================
print("\n" + "="*70)
print("CHANGES MADE:")
print("="*70)
for i, change in enumerate(changes_log, 1):
    print(f"  {i}. {change}")

print("\n" + "="*70)
print("REMAINING ISSUES FOUND:")
print("="*70)
if remaining_issues:
    for issue in remaining_issues:
        print(f"\n  Para {issue['para']}: {issue['issue']}")
        print(f"    Text: {issue['text']}")
else:
    print("  None found in main text (references may still contain citations)")

print("\n" + "="*70)
print("MANUAL ACTIONS STILL REQUIRED:")
print("="*70)
print("""
1. ADD TABLE 4 - Feature Importance Rankings
   Location: After "Table 4 presents the top 10 features ranked by XGBoost
             importance scores" in Section 5.2

   Table 4: Feature Importance Rankings
   +------+----------------------------+------------+
   | Rank | Feature                    | Importance |
   +------+----------------------------+------------+
   |  1   | amount_to_orig_balance     |   0.9770   |
   |  2   | amount                     |   0.0170   |
   |  3   | type_encoded               |   0.0031   |
   |  4   | day                        |   0.0005   |
   |  5   | hour_of_day                |   0.0005   |
   |  6   | oldbalanceOrg              |   0.0005   |
   |  7   | dest_zero_balance_before   |   0.0004   |
   |  8   | amount_to_dest_balance     |   0.0004   |
   |  9   | oldbalanceDest             |   0.0004   |
   | 10   | step                       |   0.0003   |
   +------+----------------------------+------------+

2. REFERENCES TO REVIEW:
   - Reference [5] (LightGBM paper by Ke et al.) - may no longer be cited
   - Reference [11] - check if still referenced after LightGBM removal
   - Reference [12] (Neural network paper) - no longer cited, consider removing

3. RENUMBER REFERENCES if any are removed
""")

print(f"\nCorrected document saved to:\n  {output_path}")
print("="*70)
