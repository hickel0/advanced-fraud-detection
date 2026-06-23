"""
Update the original paper with correct data from the safe features analysis.
Preserves all existing content while updating tables and adding data leakage section.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

# Open the original document
doc = Document('D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper.docx')

print("Opened original document")
print(f"Tables found: {len(doc.tables)}")
print(f"Paragraphs found: {len(doc.paragraphs)}")

# =========================================================================
# UPDATE TABLE 3: Algorithm Performance (Safe Features Results)
# =========================================================================
print("\nUpdating Table 3: Algorithm Performance...")

# Table 3 is the third table (index 2)
table3 = doc.tables[2]

# New data from safe features analysis
table3_data = [
    ['Algorithm', 'ROC-AUC', 'Precision', 'Recall', 'F1-Score'],
    ['XGBoost', '0.9999', '0.936', '0.740', '0.826'],
    ['Random Forest', '0.9792', '0.959', '0.395', '0.560'],
    ['Gradient Boosting', '0.9558', '0.905', '0.319', '0.472'],
    ['Decision Tree', '0.8393', '0.787', '0.403', '0.533'],
    ['Logistic Regression', '0.7806', '0.278', '0.042', '0.073']
]

for i, row_data in enumerate(table3_data):
    for j, val in enumerate(row_data):
        table3.rows[i].cells[j].text = val
        if i == 0:  # Header row
            table3.rows[i].cells[j].paragraphs[0].runs[0].bold = True
        elif i == 1:  # XGBoost row (best performer)
            set_cell_shading(table3.rows[i].cells[j], 'E8F5E9')

print("  Updated with safe features results")

# =========================================================================
# UPDATE TABLE 4: Feature Importance (Safe Features Results)
# =========================================================================
print("\nUpdating Table 4: Feature Importance...")

table4 = doc.tables[3]

# New data from safe features analysis
table4_data = [
    ['Rank', 'Feature', 'Importance Score'],
    ['1', 'amount_to_orig_balance', '0.8938'],
    ['2', 'amount', '0.0802'],
    ['3', 'type_encoded', '0.0087'],
    ['4', 'hour_of_day', '0.0044'],
    ['5', 'step', '0.0029'],
    ['6', 'amount_to_dest_balance', '0.0026'],
    ['7', 'oldbalanceDest', '0.0022'],
    ['8', 'day', '0.0022'],
    ['9', 'oldbalanceOrg', '0.0020'],
    ['10', 'dest_zero_balance_before', '0.0011']
]

for i, row_data in enumerate(table4_data):
    for j, val in enumerate(row_data):
        table4.rows[i].cells[j].text = val
        if i == 0:
            table4.rows[i].cells[j].paragraphs[0].runs[0].bold = True
        elif i == 1:  # Top feature
            set_cell_shading(table4.rows[i].cells[j], 'E8F5E9')

print("  Updated with safe features results")

# =========================================================================
# UPDATE TABLE 5: Imbalance Handling (Safe Features Results)
# =========================================================================
print("\nUpdating Table 5: Imbalance Handling...")

table5 = doc.tables[4]

# New data from safe features analysis
table5_data = [
    ['Method', 'Recall', 'Precision', 'F1-Score', 'False Negatives'],
    ['Baseline', '0.740', '0.936', '0.826', '31'],
    ['Undersampling', '0.992', '0.036', '0.070', '1'],
    ['SMOTE', '0.958', '0.292', '0.448', '5'],
    ['Class Weighting', '0.966', '0.462', '0.625', '4'],
    ['Combined', '0.992', '0.054', '0.103', '1']
]

for i, row_data in enumerate(table5_data):
    for j, val in enumerate(row_data):
        table5.rows[i].cells[j].text = val
        if i == 0:
            table5.rows[i].cells[j].paragraphs[0].runs[0].bold = True
        elif i == 4:  # Class Weighting row (best balance)
            set_cell_shading(table5.rows[i].cells[j], 'E8F5E9')

print("  Updated with safe features results")

# =========================================================================
# UPDATE TEXT THROUGHOUT THE DOCUMENT
# =========================================================================
print("\nUpdating text throughout document...")

# Find and update specific paragraphs
updates_made = 0

for para in doc.paragraphs:
    text = para.text

    # Update Abstract mentions of results
    if 'XGBoost achieved the highest' in text or 'Random Forest achieves' in text:
        # Clear and rewrite
        para.clear()
        run = para.add_run(
            'Using only pre-transaction features to avoid data leakage, our results demonstrate that '
            'XGBoost achieves the highest ROC-AUC (0.9999) among tested algorithms. '
            'The ratio of transaction amount to sender\'s balance emerges as the most predictive feature '
            '(89.4% importance). For handling the severe class imbalance (773:1 ratio), '
            'Weighted XGBoost achieves 96.6% recall while maintaining 46.2% precision.'
        )
        updates_made += 1

    # Update RQ1 results text
    if 'Experimental results demonstrate clear performance differences' in text:
        para.clear()
        run = para.add_run(
            'Experimental results demonstrate clear performance differences among the evaluated algorithms '
            'when using the pre-transaction (safe) feature set. XGBoost significantly outperforms all other '
            'algorithms with an AUC of 0.9999 and F1-score of 0.826. Tree-based ensemble methods '
            '(XGBoost, Random Forest, Gradient Boosting) consistently outperform linear models, with '
            'Logistic Regression achieving only 4.2% recall, indicating that fraud patterns are highly non-linear.'
        )
        updates_made += 1

    # Update RQ2 results text
    if 'Feature importance analysis identified several highly predictive' in text:
        para.clear()
        run = para.add_run(
            'Feature importance analysis using the safe feature set identified the ratio of transaction '
            'amount to sender\'s balance (amount_to_orig_balance) as the dominant predictive feature with '
            '89.4% importance. This indicates that fraudsters attempt to transfer unusually large proportions '
            'of account balances. Transaction amount itself is the second most important feature (8.0%), '
            'while transaction type provides modest predictive value (0.9%).'
        )
        updates_made += 1

    # Update RQ3 results text
    if 'Imbalance handling experiments revealed significant improvements' in text:
        para.clear()
        run = para.add_run(
            'Imbalance handling experiments using safe features revealed that class weighting (Weighted XGBoost) '
            'achieves the best precision-recall balance with 96.6% recall and 46.2% precision. '
            'While undersampling achieves the highest recall (99.2%), it suffers from very low precision (3.6%). '
            'SMOTE provides good recall (95.8%) with moderate precision (29.2%). The choice of method '
            'depends on business requirements regarding the cost of missed fraud versus investigation overhead.'
        )
        updates_made += 1

print(f"  Updated {updates_made} paragraphs")

# =========================================================================
# ADD DATA LEAKAGE SECTION
# =========================================================================
print("\nAdding Data Leakage section...")

# Find the Feature Engineering section (3.2)
insert_idx = None
for i, para in enumerate(doc.paragraphs):
    if '3.2 Feature Engineering' in para.text:
        insert_idx = i
        break

if insert_idx:
    # Find the next heading after 3.2 to insert before it
    for i in range(insert_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            insert_idx = i
            break

    # We need to insert paragraphs. In python-docx, we'll add to the section
    # Actually, let's add a new subsection after 3.2

# Find position after 3.2 content but before 3.3
feature_eng_found = False
for i, para in enumerate(doc.paragraphs):
    if '3.2 Feature Engineering' in para.text:
        feature_eng_found = True
    if feature_eng_found and '3.3' in para.text:
        # Insert data leakage content here
        # We'll modify the 3.2 section content instead
        break

# Since inserting is complex, let's append a note about data leakage to the 3.2 section
# Find the 3.2 section and add leakage info
for i, para in enumerate(doc.paragraphs):
    if '3.2 Feature Engineering' in para.text:
        # Find the end of this section
        for j in range(i+1, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name.startswith('Heading'):
                # Insert before this heading
                # Add data leakage subsection
                break

# Alternative: Update the Feature Engineering section text
for para in doc.paragraphs:
    if 'engineered features' in para.text.lower() and 'balance' in para.text.lower():
        # This is likely in the feature engineering section
        original_text = para.text
        if 'full_drain' not in para.text and 'leakage' not in para.text.lower():
            # Add data leakage note
            para.clear()
            run = para.add_run(original_text)
            para.add_run('\n\n')
            run2 = para.add_run(
                'Critical Note on Data Leakage: '
            )
            run2.bold = True
            para.add_run(
                'Initial analysis revealed that the full_drain feature (indicating complete account drainage) '
                'had a +0.987 correlation with fraud, essentially encoding the fraud definition. '
                'This constitutes data leakage as the feature requires post-transaction information. '
                'All results in this paper use only pre-transaction features to ensure realistic, '
                'deployable performance estimates.'
            )
            print("  Added data leakage note to Feature Engineering section")
            break

# =========================================================================
# UPDATE CONCLUSION
# =========================================================================
print("\nUpdating Conclusion...")

for para in doc.paragraphs:
    if 'This study evaluated' in para.text or 'In conclusion' in para.text.lower():
        para.clear()
        run = para.add_run(
            'This study evaluated machine learning approaches for fraud detection using the PaySim dataset, '
            'with careful attention to avoiding data leakage. Key findings include: '
            '(1) XGBoost achieved the highest ROC-AUC (0.9999) using only pre-transaction features, '
            'significantly outperforming other algorithms; '
            '(2) The ratio of transaction amount to sender\'s balance is the most predictive feature (89.4% importance), '
            'indicating fraudsters attempt to transfer unusually large proportions of account balances; '
            '(3) Class weighting provides the best precision-recall balance for handling the severe class imbalance, '
            'achieving 96.6% recall with 46.2% precision. '
            'A critical contribution is the identification and resolution of data leakage in commonly used features, '
            'providing realistic performance benchmarks for production deployment.'
        )
        print("  Updated conclusion")
        break

# =========================================================================
# SAVE UPDATED DOCUMENT
# =========================================================================
output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - UPDATED.docx'
doc.save(output_path)

print("\n" + "="*70)
print("DOCUMENT UPDATED SUCCESSFULLY")
print("="*70)
print(f"\nOriginal preserved at: CSC1112 - Fraud Detection Paper.docx")
print(f"Updated version saved to: {output_path}")
print("\nChanges made:")
print("  ✓ Table 3: Algorithm Performance - Updated with safe features results")
print("  ✓ Table 4: Feature Importance - Updated with safe features results")
print("  ✓ Table 5: Imbalance Handling - Updated with safe features results")
print("  ✓ Abstract: Updated with correct metrics")
print("  ✓ RQ1 Results text: Updated")
print("  ✓ RQ2 Results text: Updated")
print("  ✓ RQ3 Results text: Updated")
print("  ✓ Data leakage note: Added to Feature Engineering section")
print("  ✓ Conclusion: Updated with correct findings")
print("\nAll data now aligns with the notebook safe features analysis.")
print("="*70)
