"""
Final comprehensive fix for the Fraud Detection Paper.
Removes all LightGBM/neural network references and adds Table 4.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Load the document
doc = Document('D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper (1).docx')

changes_log = []

# ============================================================================
# Helper function to replace text while preserving some formatting
# ============================================================================
def replace_text_in_paragraph(para, replacements):
    """Apply multiple replacements to a paragraph."""
    if not para.text.strip():
        return False

    original_text = para.text
    new_text = original_text

    for old, new in replacements:
        if old in new_text:
            new_text = new_text.replace(old, new)

    if new_text != original_text:
        # Clear runs and set new text
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        return True
    return False

# ============================================================================
# Define all replacements
# ============================================================================
replacements = [
    # Introduction - algorithm list (multiple variants)
    (
        "Several models are compared, including logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks.",
        "Several models are compared, including Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost."
    ),
    (
        "logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks",
        "Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost"
    ),
    (
        "XGBoost, LightGBM, and neural networks",
        "and XGBoost"
    ),

    # Related Work - neural network sentence
    (
        "While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12]. ",
        ""
    ),
    (
        "While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12].",
        ""
    ),

    # Related Work - ensemble approaches
    (
        "including ensemble approaches and neural networks, to determine their effectiveness",
        "including tree-based ensemble methods, to determine their effectiveness"
    ),
    (
        "including ensemble approaches and neural networks",
        "including tree-based ensemble methods"
    ),

    # Related Work - LightGBM in gradient boosting discussion
    (
        "More recent studies further highlight the effectiveness of gradient boosting techniques such as XGBoost and LightGBM, which consistently achieve strong performance on structured financial datasets due to their scalability and ability to model complex feature interactions [10], [11].",
        "More recent studies further highlight the effectiveness of gradient boosting techniques, particularly XGBoost, which consistently achieves strong performance on structured financial datasets due to its scalability and ability to model complex feature interactions [10]."
    ),
    (
        "such as XGBoost and LightGBM",
        ", particularly XGBoost,"
    ),
    (
        "[10], [11]",
        "[10]"
    ),
]

# Apply replacements to all paragraphs
for para in doc.paragraphs:
    for old, new in replacements:
        if old in para.text:
            original = para.text
            new_text = para.text.replace(old, new)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            changes_log.append(f"Replaced: '{old[:50]}...' -> '{new[:50]}...'")

# ============================================================================
# Find the location to insert Table 4 and add it
# ============================================================================
table4_inserted = False
insert_after_idx = None

for i, para in enumerate(doc.paragraphs):
    if "Table 4 presents the top 10 features" in para.text or "Table 4 presents the top" in para.text:
        insert_after_idx = i
        break

if insert_after_idx is not None:
    # Find the paragraph element
    para_element = doc.paragraphs[insert_after_idx]._element

    # Create Table 4
    table = doc.add_table(rows=11, cols=3)
    # Try to set table style, use default if not available
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # Use default style

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Rank'
    header_cells[1].text = 'Feature'
    header_cells[2].text = 'Importance'

    # Make header bold
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    data = [
        ('1', 'amount_to_orig_balance', '0.9770'),
        ('2', 'amount', '0.0170'),
        ('3', 'type_encoded', '0.0031'),
        ('4', 'day', '0.0005'),
        ('5', 'hour_of_day', '0.0005'),
        ('6', 'oldbalanceOrg', '0.0005'),
        ('7', 'dest_zero_balance_before', '0.0004'),
        ('8', 'amount_to_dest_balance', '0.0004'),
        ('9', 'oldbalanceDest', '0.0004'),
        ('10', 'step', '0.0003'),
    ]

    for row_idx, (rank, feature, importance) in enumerate(data, 1):
        row = table.rows[row_idx]
        row.cells[0].text = rank
        row.cells[1].text = feature
        row.cells[2].text = importance

    # Move table to correct position (after the paragraph that mentions it)
    # This is tricky in python-docx, so we'll note it needs manual adjustment
    table4_inserted = True
    changes_log.append("Added Table 4 (Feature Importance Rankings) - may need manual repositioning")

# ============================================================================
# Final scan for remaining issues
# ============================================================================
print("="*70)
print("FINAL SCAN FOR REMAINING ISSUES...")
print("="*70)

remaining_issues = []
references_to_check = []

for i, para in enumerate(doc.paragraphs):
    text = para.text
    text_lower = text.lower()

    # Skip empty
    if not text.strip():
        continue

    # Check for LightGBM (excluding references section - check for year patterns)
    if 'lightgbm' in text_lower:
        if '2017' not in text and 'ke, g' not in text_lower:
            remaining_issues.append(f"Para {i+1}: LightGBM found: {text[:100]}...")
        else:
            references_to_check.append(f"Reference with LightGBM: {text[:80]}...")

    # Check for neural network (excluding references and graph neural networks)
    if 'neural network' in text_lower:
        if 'graph neural' not in text_lower and '2021' not in text and 'rb, a' not in text_lower:
            remaining_issues.append(f"Para {i+1}: Neural network found: {text[:100]}...")

# ============================================================================
# Save
# ============================================================================
output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - CORRECTED.docx'
doc.save(output_path)

# ============================================================================
# Report
# ============================================================================
print("\n" + "="*70)
print("CHANGES MADE:")
print("="*70)
for change in changes_log:
    print(f"  - {change}")

print("\n" + "="*70)
print("REMAINING ISSUES IN MAIN TEXT:")
print("="*70)
if remaining_issues:
    for issue in remaining_issues:
        print(f"  ! {issue}")
else:
    print("  None - all LightGBM/neural network references removed from main text")

print("\n" + "="*70)
print("REFERENCES SECTION - MANUAL REVIEW NEEDED:")
print("="*70)
print("""
The following references may no longer be cited and should be reviewed:

  [5]  Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting
       decision tree.
       STATUS: LightGBM no longer discussed - CONSIDER REMOVING

  [11] (Check what this references - may have been for LightGBM)
       STATUS: Citation [11] was removed from text - CONSIDER REMOVING

  [12] RB, A., & KR, S.K. (2021). Credit card fraud detection using artificial
       neural network.
       STATUS: Neural networks no longer discussed - CONSIDER REMOVING

If you remove references, remember to RENUMBER all subsequent references
and update all citations in the text.
""")

print("="*70)
print("TABLE 4 STATUS:")
print("="*70)
if table4_inserted:
    print("""
  Table 4 was added to the document but appears at the END of the document.
  You need to MANUALLY MOVE IT to the correct location:
    - After the sentence "Table 4 presents the top 10 features ranked by
      XGBoost importance scores."
    - In Section 5.2 (Feature Importance)
""")
else:
    print("  Could not find location to insert Table 4 - add manually")

print(f"\nCorrected document saved to:\n  {output_path}")
print("="*70)
