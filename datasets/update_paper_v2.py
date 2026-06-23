"""
Script to update the Fraud Detection Paper with Option 3:
- Present both results (with and without full_drain)
- Frame the finding about data leakage as a discovery
- Add limitations discussion
All changes are highlighted in yellow for review.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_updated_paper_v2():
    """Create the updated paper with Option 3 implementation"""
    doc = Document()

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('Fraud Detection in Financial Transactions Using Machine Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('An Analysis of the PaySim Synthetic Financial Dataset')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Lee Hickey | Emma Reen | Sarah O\'Hanlon\n').bold = True
    authors.add_run('lee.hickey28@mail.dcu.ie | emma.reen2@mail.dcu.ie | sarah.ohanlon24@mail.dcu.ie\n')
    authors.add_run('21407166 | 21491492 | 22444924\n')
    authors.add_run('School of Computing, Faculty of Engineering and Computing\n')
    authors.add_run('Dublin City University')

    doc.add_paragraph()

    # =========================================================================
    # ABSTRACT - UPDATED
    # =========================================================================
    doc.add_heading('Abstract', level=1)

    abstract = doc.add_paragraph()
    abstract.add_run('Financial fraud poses a significant threat to the integrity of digital payment systems, resulting in substantial economic losses worldwide. This study investigates the application of machine learning techniques for detecting fraudulent transactions in mobile money systems. Using the PaySim synthetic financial dataset, which simulates real-world mobile money transactions, we develop and evaluate classification models capable of accurately identifying fraudulent activities. The dataset contains over 6.3 million transactions across five transaction types, with approximately 8,213 labeled fraudulent cases representing a highly imbalanced classification problem (0.13% fraud rate). Our research explores various approaches to handle class imbalance, feature engineering strategies based on transaction patterns and account behaviors, and the comparative performance of different machine learning algorithms for fraud detection. Results demonstrate that ')

    run = abstract.add_run('Random Forest achieves the highest ROC-AUC (0.999)')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    abstract.add_run(' when using engineered features. ')

    # NEW: Add critical finding to abstract
    run = abstract.add_run('Critically, our analysis reveals that a single engineered feature—complete account drainage (full_drain)—captures 97.5% of fraudulent transactions, suggesting that fraud in PaySim follows a deterministic pattern. When this feature is excluded to simulate more realistic conditions, model recall drops from 99% to 69%, providing important insights into the gap between synthetic and real-world fraud detection performance.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('(estimated 750 words, 15 marks) EMMA')

    intro_text = """The growth of digital financial services has transformed how individuals and businesses transfer money. Mobile banking, digital wallets, and online payment systems now process millions of transactions each day improving convenience, speed, and accessibility. However, this expansion has also increased the risk of financial fraud. Fraudulent transactions can cause financial losses, reduce customer trust, and create legal and reputational problems for the service providers. As a result, fraud detection has become an important application of data science and artificial intelligence in finance.

Traditional fraud detection systems have often relied on rule-based methods. These systems flag transactions based on fixed conditions, such as unusually large amounts or repeated transfers. Although this approach can detect simple and known fraud patterns, it is often ineffective against more complex or changing behaviour. Fraudsters can adapt their actions to avoid these predefined rules, which limits the effectiveness of static systems. Machine learning offers a more flexible alternative because it can learn patterns directly from historical transaction data. Instead of depending only on manually created rules, machine learning models can detect complex relationships between features and identify suspicious behaviour more effectively [1].

Despite its advantages, fraud detection remains a difficult machine learning task. One major challenge is class imbalance. In most transactions datasets, fraudulent cases account for only a very small proportion of all records. This means a classifier can achieve high overall accuracy by predicting nearly all transactions as legitimate, while still failing to identify the fraud cases that matter most. For this reason evaluation metrics such as precision, recall, F1-score, and ROC-AUC are more useful than accuracy alone when assessing fraud detection models [2]. In real-world applications, false negatives are particularly costly as they allow fraudulent transactions to go undetected.

Another challenge is that fraudulent behaviour is often hidden within normal transaction activity. Raw features such as amount, transaction type, or account balance may not be sufficient on their own to identify suspicious behaviour. Therefore, feature engineering is often necessary to improve model performance. Features based on balance changes, account draining, transaction ratios, and deviations from normal account behaviour may reveal useful fraud patterns.

Fraud detection is also not a static problem. Transaction behaviour and fraud strategies can change over time, meaning a model trained on past data may lose effectiveness if it is not regularly evaluated and updated.

This paper investigates fraud detection using the PaySim synthetic financial dataset. PaySim was created to simulate mobile money transactions using aggregated characteristics from real financial logs, while avoiding the privacy issues associated with real customer data [1]. The dataset contains 6,362,620 transactions and 11 original features, with only 8,213 cases labelled as fraudulent, representing approximately 0.13% of the total data [3]. It includes five transaction types: CASH_IN, CASH_OUT, DEBIT, PAYMENT, and TRANSFER. The dataset also records sender and receiver balances before and after each transaction, making it useful for creating behavioural and balance-related features. Due to its large size and strong imbalance, PaySim is well suited to studying fraud detection methods.

The main aim of this study is to evaluate how effectively machine learning algorithms can detect fraudulent financial transactions in a highly imbalanced dataset. Several models are compared, including logistic regression, decision trees, random forests, gradient boosting, and XGBoost. Comparing these algorithms helps determine which methods are most effective for structured financial transaction data. Previous research has shown that tree-based ensemble methods such as XGBoost and Random Forest often perform strongly in tabular classification tasks because they can capture non-linear relationships between variables [4][5].

In addition to comparing algorithms, this paper investigates both the features that are most predictive of fraudulent behaviour and the methods used to address class imbalance. Feature importance analysis is valuable not only for improving predictive performance but also for supporting interpretability, which is particularly important in financial applications where organisations may need to explain why specific transactions are flagged as suspicious. Features such as full account drainage, unusual transaction size, and abnormal balance changes may serve as strong indicators of fraud. Furthermore, because the fraud class is severely underrepresented, the study evaluates imbalance handling techniques including undersampling, oversampling with SMOTE, and class weighting. SMOTE is widely used as it generates synthetic minority examples and can improve learning for rare classes [6].

Overall, this paper addresses three research questions:
1. Which machine learning algorithms perform best for fraud detection in imbalanced transaction data?
2. Which features are most predictive of fraudulent activity?
3. How can class imbalance be handled effectively?

By answering these questions, the study aims to show how AI and data science can support more reliable fraud detection in digital payment systems."""

    doc.add_paragraph(intro_text)

    # =========================================================================
    # RELATED WORK
    # =========================================================================
    doc.add_heading('2. Related Work', level=1)
    doc.add_paragraph('(estimated 750 words, 15 marks) SARAH')

    related_work_text = """Financial fraud detection has become a widely studied problem due to the rapid growth of digital and mobile transactions. Machine learning techniques are increasingly used to identify fraudulent behaviour by analysing transaction patterns at scale. However, fraud detection presents unique challenges, particularly due to its temporal nature and the risk of unrealistic evaluation. Dal Pozzolo et al. highlight that many traditional approaches rely on random data splits, which can lead to overly optimistic results and fail to reflect real-world conditions [7]. In addition, Dornadula and Geetha emphasise that fraud detection systems benefit from modelling behavioural patterns in transaction streams, rather than treating transactions independently [8].

A significant body of research has focused on comparing machine learning algorithms for fraud detection. Sailusha et al. demonstrate that ensemble methods, particularly Random Forest, outperform simpler models due to their ability to capture complex, nonlinear relationships in transaction data [9]. More recent studies further highlight the effectiveness of gradient boosting techniques such as XGBoost and LightGBM, which consistently achieve strong performance on structured financial datasets due to their scalability and ability to model complex feature interactions [10], [11]. While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12]. These findings align with this study, which evaluates a range of machine learning models, including ensemble approaches, to determine their effectiveness in detecting fraudulent transactions.

A fundamental challenge in fraud detection is the severe class imbalance present in financial transaction datasets, where fraudulent transactions typically represent a very small proportion of the data. This imbalance can significantly bias machine learning models towards the majority class, resulting in high overall accuracy but poor detection of fraudulent cases. As highlighted by Rubaidi et al., traditional evaluation metrics such as accuracy are insufficient in such contexts, and greater emphasis must be placed on metrics such as precision, recall, and F1-score [13]. Various techniques, including oversampling, undersampling, and cost-sensitive learning, have been proposed to address this issue. Effectively handling class imbalance is therefore essential for developing reliable fraud detection systems.

Another critical aspect of fraud detection is the temporal nature of transaction data and the presence of concept drift. As fraud patterns evolve over time, models trained on historical data may become less effective when applied to future transactions. Dal Pozzolo et al. emphasise that chronological evaluation is essential to avoid data leakage and to ensure realistic performance assessment [7]. In addition, recent studies highlight that concept drift can significantly impact model reliability, as changes in user behaviour and fraud strategies lead to shifts in data distributions over time [14]. To address this, techniques such as sliding window evaluation, adaptive learning, and drift detection methods have been proposed to monitor and maintain model performance. These approaches highlight the importance of incorporating temporal analysis into fraud detection systems.

Explainability has become an important consideration in fraud detection systems, particularly in financial domains where model transparency is required for trust and regulatory compliance. Methods such as SHAP (SHapley Additive exPlanations) provide a framework for interpreting complex machine learning models by quantifying the contribution of each feature to a prediction. Recent studies demonstrate that SHAP enables both global and local interpretability, allowing practitioners to identify key factors influencing fraud predictions and to explain individual decisions in a consistent manner [15]. Furthermore, explainable AI techniques have been shown to improve the usability and reliability of fraud detection systems by supporting human decision-making and facilitating model validation [16]. These approaches are particularly valuable when using high-performing but complex models such as ensemble methods, where interpretability would otherwise be limited.

Overall, existing research highlights the importance of behavioural modelling, appropriate handling of class imbalance, and the use of advanced machine learning techniques for fraud detection. In addition, temporal evaluation and concept drift analysis are essential for ensuring realistic and robust model performance. These findings collectively inform the methodology adopted in this study, which integrates these approaches to provide a comprehensive evaluation of fraud detection techniques."""

    doc.add_paragraph(related_work_text)

    # =========================================================================
    # METHODOLOGY
    # =========================================================================
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph('(estimated 500 words, 30 marks) LEE')

    doc.add_heading('3.1 Dataset Description', level=2)

    dataset_text = """This study utilizes the PaySim synthetic financial dataset, comprising 6,362,620 transactions simulating one month of mobile money activity. The dataset includes five transaction types: CASH_IN (deposits), CASH_OUT (withdrawals), DEBIT (direct debits), PAYMENT (merchant payments), and TRANSFER (peer-to-peer transfers). Table 1 summarizes the dataset characteristics."""
    doc.add_paragraph(dataset_text)

    # Table 1 - Dataset Statistics
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    table1.rows[0].cells[0].text = 'Attribute'
    table1.rows[0].cells[1].text = 'Details'
    table1.rows[1].cells[0].text = 'Total Records'
    table1.rows[1].cells[1].text = '6,362,620 transactions'
    table1.rows[2].cells[0].text = 'Features'
    table1.rows[2].cells[1].text = '11 columns'
    table1.rows[3].cells[0].text = 'Fraudulent Cases'
    table1.rows[3].cells[1].text = '8,213 (0.13%)'
    table1.rows[4].cells[0].text = 'Transaction Types'
    table1.rows[4].cells[1].text = 'CASH_IN, CASH_OUT, DEBIT, PAYMENT, TRANSFER'
    doc.add_paragraph('Table 1: PaySim Dataset Statistics')

    dataset_text2 = """Each transaction record contains 11 features: step (hourly time unit), type (transaction category), amount (transaction value), nameOrig (originating account identifier), oldbalanceOrg and newbalanceOrig (account balance before and after transaction), nameDest (destination account identifier), oldbalanceDest and newbalanceDest (destination balance before and after), isFraud (target variable indicating fraudulent transactions), and isFlaggedFraud (system-generated fraud flag). The target variable, isFraud, contains 8,213 positive cases (0.13%), representing an extreme class imbalance ratio of approximately 773:1."""
    doc.add_paragraph(dataset_text2)

    doc.add_heading('3.2 Feature Engineering', level=2)

    fe_text = """We developed engineered features capturing transaction patterns and account behaviors indicative of fraud, informed by domain knowledge and prior literature. The engineered features include:

• Balance Change Features: Computed as the difference between pre- and post-transaction balances for both originating (balance_change_orig) and destination (balance_change_dest) accounts.

• Error Flags: Binary indicators identifying discrepancies between expected and actual balance changes (orig_balance_error, dest_balance_error), which may indicate fraudulent manipulation of account records.

• Zero Balance Indicators: Flags for transactions resulting in zero balance (orig_zero_balance) or originating from zero-balance destination accounts (dest_zero_balance_before).

• Amount-to-Balance Ratios: Normalized transaction amounts relative to account balances (amount_to_orig_balance, amount_to_dest_balance), capturing unusually large transactions relative to account capacity.

• Full Drainage Flag: A binary indicator specifically identifying transactions that completely empty an account where the transaction amount equals the original balance.

• Temporal Features: Hour-of-day and day-of-simulation extracted from the step variable to capture time-based fraud patterns.

Categorical encoding transformed the transaction type variable into numeric format using label encoding. The final feature set comprises 18 features, including both original and engineered variables. All numerical features were standardized using z-score normalization to ensure comparable scales for algorithms sensitive to feature magnitudes. Infinite values resulting from division operations were replaced with zeros."""
    doc.add_paragraph(fe_text)

    doc.add_heading('3.3 Experimental Design', level=2)

    exp_design = """Data was partitioned using stratified sampling to preserve class distributions, with 70% allocated to training and 30% to testing. Stratification ensures that both training and test sets maintain the same fraud ratio as the original dataset, enabling unbiased performance evaluation. To manage computational requirements while maintaining representative samples, model comparison experiments utilized a random sample of 500,000 training instances.

For Research Question 1 (algorithm comparison), we evaluated five classification algorithms: Logistic Regression serving as a baseline linear model, Decision Tree as an interpretable single-tree classifier, Random Forest implementing bagging ensemble methodology, Gradient Boosting as a sequential boosting ensemble, and XGBoost as an optimized gradient boosting implementation. All algorithms were implemented using scikit-learn and XGBoost libraries.

For Research Question 2 (feature importance), we extracted importance scores from the best-performing model using built-in feature importance measures. Correlation analysis computed Pearson coefficients between each feature and the target variable.

For Research Question 3 (class imbalance handling), we evaluated five approaches: baseline with no handling, random undersampling reducing majority class to minority size, SMOTE oversampling generating synthetic minority samples, class weight adjustment using scale_pos_weight parameter, and a combined approach using partial undersampling with class weights. All resampling methods were implemented using the imbalanced-learn library."""
    doc.add_paragraph(exp_design)

    doc.add_heading('3.4 Evaluation Metrics', level=2)

    metrics_text = """Given the severe class imbalance, we employed multiple evaluation metrics beyond simple accuracy. ROC-AUC (Area Under the Receiver Operating Characteristic Curve) measures discrimination ability across all classification thresholds. Precision quantifies the proportion of predicted frauds that are actually fraudulent. Recall (sensitivity) measures the proportion of actual frauds correctly identified. F1-score provides the harmonic mean of precision and recall. Confusion matrix analysis provides counts of true positives, false negatives (missed frauds), false positives (false alarms), and true negatives."""
    doc.add_paragraph(metrics_text)

    # =========================================================================
    # EXPERIMENTS
    # =========================================================================
    doc.add_heading('4. Experiments', level=1)
    doc.add_paragraph('(estimated 500 words, 20 marks) LEE')

    doc.add_heading('4.1 Algorithm Comparison (RQ1)', level=2)

    exp_text = """We trained five classification algorithms on the standardized feature set to determine optimal algorithm selection for fraud detection. Each algorithm was configured within a scikit-learn Pipeline incorporating StandardScaler preprocessing to ensure consistent feature scaling. Table 2 presents the hyperparameter configurations for each algorithm."""
    doc.add_paragraph(exp_text)

    # Table 2 - Hyperparameters
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = 'Algorithm'
    table2.rows[0].cells[1].text = 'Configuration'
    table2.rows[1].cells[0].text = 'Logistic Regression'
    table2.rows[1].cells[1].text = 'L2 regularization, max_iter=1000'
    table2.rows[2].cells[0].text = 'Decision Tree'
    table2.rows[2].cells[1].text = 'max_depth=10'
    table2.rows[3].cells[0].text = 'Random Forest'
    table2.rows[3].cells[1].text = 'n_estimators=100, max_depth=10, n_jobs=-1'
    table2.rows[4].cells[0].text = 'Gradient Boosting'
    table2.rows[4].cells[1].text = 'n_estimators=100, max_depth=4, learning_rate=0.1'
    table2.rows[5].cells[0].text = 'XGBoost'
    table2.rows[5].cells[1].text = 'n_estimators=100, max_depth=4, learning_rate=0.1, eval_metric=logloss'
    doc.add_paragraph('Table 2: Algorithm Hyperparameter Configurations')

    exp_text2 = """Models were trained on the 500,000-instance training sample and evaluated on the complete test set comprising approximately 1.9 million transactions. We computed ROC-AUC, accuracy, precision, recall, F1-score, and average precision for each algorithm."""
    doc.add_paragraph(exp_text2)

    doc.add_heading('4.2 Feature Importance Analysis (RQ2)', level=2)

    fi_text = """Using XGBoost with increased capacity (n_estimators=200, max_depth=6) for feature importance analysis, incorporating scale_pos_weight equal to the imbalance ratio. Three complementary analysis approaches were employed: built-in XGBoost feature importance scores based on information gain, Pearson correlation coefficients between features and the target variable, and distribution comparisons examining statistical differences between fraudulent and legitimate transaction groups."""
    doc.add_paragraph(fi_text)

    doc.add_heading('4.3 Imbalance Handling Evaluation (RQ3)', level=2)

    imb_text = """Five imbalance handling strategies were implemented using imbalanced-learn pipelines with XGBoost as the base classifier. The baseline approach applied no resampling. Random undersampling reduced the majority class to match minority class size. SMOTE generated synthetic minority samples using k=5 nearest neighbors. Class weighting set scale_pos_weight equal to the imbalance ratio (~773). The combined approach applied undersampling to achieve a 2:1 ratio followed by class weighting of 2."""
    doc.add_paragraph(imb_text)

    # =========================================================================
    # RESULTS - WITH BOTH ANALYSES
    # =========================================================================
    doc.add_heading('5. Results', level=1)
    doc.add_paragraph('(estimated 500 words, 15 marks) LEE')

    doc.add_heading('5.1 Algorithm Performance (RQ1)', level=2)

    results_intro = doc.add_paragraph()
    results_intro.add_run('Experimental results demonstrate exceptional performance across all algorithms when using the full engineered feature set. Table 3 presents the comprehensive performance metrics.')

    # Table 3 - Algorithm Performance (WITH full_drain)
    table3 = doc.add_table(rows=6, cols=5)
    table3.style = 'Table Grid'
    headers = ['Algorithm', 'ROC-AUC', 'Precision', 'Recall', 'F1-Score']
    for i, header in enumerate(headers):
        table3.rows[0].cells[i].text = header

    data = [
        ['Logistic Regression', '0.9980', '0.9984', '0.9935', '0.9959'],
        ['Decision Tree', '0.9968', '1.0000', '0.9935', '0.9967'],
        ['Random Forest', '0.9992', '1.0000', '0.9947', '0.9974'],
        ['Gradient Boosting', '0.9968', '1.0000', '0.9935', '0.9967'],
        ['XGBoost', '0.9972', '0.9988', '0.9939', '0.9963'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table3.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph('Table 3: Algorithm Performance with Full Feature Set')

    results_disc = doc.add_paragraph()
    results_disc.add_run('Random Forest achieved the highest ROC-AUC (0.9992), with all ensemble methods achieving AUC above 0.996. Notably, even Logistic Regression achieved 0.998 AUC, suggesting strong linear separability in the feature space.')

    # =========================================================================
    # NEW SECTION: Critical Finding on full_drain
    # =========================================================================
    doc.add_heading('5.2 Critical Finding: Feature Dominance Analysis', level=2)

    critical_note = doc.add_paragraph()
    run = critical_note.add_run('[NEW SECTION - Critical Analysis of Results]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    finding_intro = doc.add_paragraph()
    run = finding_intro.add_run('Investigation of the exceptionally high performance metrics revealed a critical finding regarding the full_drain feature. Further analysis was conducted to understand why all models achieved near-perfect scores.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Explain the finding
    finding_text = doc.add_paragraph()
    run = finding_text.add_run('The full_drain feature, which identifies transactions that completely empty an account (where amount equals oldbalanceOrg and newbalanceOrig equals zero), was found to have an extremely strong relationship with the fraud label:')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Add bullet points for the finding
    bullets = doc.add_paragraph()
    run = bullets.add_run('''
• 100% of transactions with full_drain=1 are fraudulent (8,008 transactions)
• 97.5% of all fraudulent transactions (8,008 of 8,213) have full_drain=1
• The correlation between full_drain and isFraud is +0.987 (near-perfect)
• Using only full_drain as a prediction rule achieves 97.5% recall with 100% precision''')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    interpretation = doc.add_paragraph()
    run = interpretation.add_run('This finding indicates that fraud in the PaySim dataset is defined by a deterministic pattern: fraudulent transactions are specifically those that completely drain victim accounts. The full_drain feature essentially encodes this fraud definition, explaining both its 92.6% feature importance and the exceptional model performance.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # NEW: Comparative Analysis WITHOUT full_drain
    # =========================================================================
    doc.add_heading('5.3 Realistic Performance Estimation', level=2)

    realistic_intro = doc.add_paragraph()
    run = realistic_intro.add_run('To estimate more realistic performance that might be expected on real-world data where fraud patterns are less deterministic, we re-evaluated the models after excluding the full_drain feature. Table 4 presents this comparative analysis.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Table 4 - Comparison WITH vs WITHOUT full_drain
    compare_note = doc.add_paragraph()
    run = compare_note.add_run('[NEW TABLE - Comparing performance with and without full_drain feature]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    table4 = doc.add_table(rows=3, cols=5)
    table4.style = 'Table Grid'

    # Headers
    table4.rows[0].cells[0].text = 'Scenario'
    table4.rows[0].cells[1].text = 'ROC-AUC'
    table4.rows[0].cells[2].text = 'Precision'
    table4.rows[0].cells[3].text = 'Recall'
    table4.rows[0].cells[4].text = 'F1-Score'

    # With full_drain
    table4.rows[1].cells[0].text = 'With full_drain (optimistic)'
    table4.rows[1].cells[1].text = '0.9972'
    table4.rows[1].cells[2].text = '0.9988'
    table4.rows[1].cells[3].text = '0.9939'
    table4.rows[1].cells[4].text = '0.9963'

    # Without full_drain - highlight this row
    table4.rows[2].cells[0].text = 'Without full_drain (realistic)'
    table4.rows[2].cells[1].text = '0.9962'
    table4.rows[2].cells[2].text = '0.9198'
    table4.rows[2].cells[3].text = '0.6887'
    table4.rows[2].cells[4].text = '0.7877'

    # Highlight the "without" row
    for cell in table4.rows[2].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_paragraph('Table 4: XGBoost Performance Comparison (With vs Without full_drain)')

    # Discussion of realistic results
    realistic_disc = doc.add_paragraph()
    run = realistic_disc.add_run('Without the full_drain feature, model recall drops significantly from 99.4% to 68.9%, representing a 31 percentage point decrease. This finding has important implications:')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    implications = doc.add_paragraph()
    run = implications.add_run('''
1. The ~99% recall achieved with full_drain reflects the synthetic nature of PaySim's fraud definition rather than the model's ability to detect complex fraud patterns.

2. The ~69% recall without full_drain provides a more realistic estimate of performance that might be expected on real-world data where fraud patterns are diverse and less predictable.

3. The high AUC (0.996) is maintained even without full_drain, indicating the model still achieves good discrimination, but the practical recall at standard thresholds is substantially lower.''')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # 5.4 Feature Importance (Updated)
    # =========================================================================
    doc.add_heading('5.4 Feature Importance (RQ2)', level=2)

    fi_intro = doc.add_paragraph()
    fi_intro.add_run('Feature importance analysis confirmed the dominance of the full_drain feature. Table 5 presents the top features ranked by XGBoost importance scores.')

    table5 = doc.add_table(rows=6, cols=3)
    table5.style = 'Table Grid'
    table5.rows[0].cells[0].text = 'Rank'
    table5.rows[0].cells[1].text = 'Feature'
    table5.rows[0].cells[2].text = 'Importance'

    feature_data = [
        ['1', 'full_drain', '0.9264'],
        ['2', 'amount_to_dest_balance', '0.0731'],
        ['3', 'oldbalanceOrg', '0.0005'],
        ['4', 'balance_change_orig', '0.0001'],
        ['5', 'Other features', '<0.0001'],
    ]

    for row_idx, row_data in enumerate(feature_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table5.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph('Table 5: Feature Importance Ranking')

    fi_disc = doc.add_paragraph()
    fi_disc.add_run('The full_drain feature accounts for 92.64% of total importance, with amount_to_dest_balance contributing 7.31%. Together, these two features account for 99.95% of the model\'s predictive power, with all other features contributing negligibly.')

    # =========================================================================
    # 5.5 Imbalance Handling (RQ3)
    # =========================================================================
    doc.add_heading('5.5 Imbalance Handling (RQ3)', level=2)

    imb_intro = doc.add_paragraph()
    imb_intro.add_run('Table 6 presents the comparison of imbalance handling methods using the full feature set.')

    table6 = doc.add_table(rows=6, cols=5)
    table6.style = 'Table Grid'
    headers6 = ['Method', 'Recall', 'Precision', 'F1-Score', 'False Negatives']
    for i, header in enumerate(headers6):
        table6.rows[0].cells[i].text = header

    imb_data = [
        ['Baseline', '0.9939', '0.9988', '0.9964', '15'],
        ['Undersampling', '0.9943', '0.3923', '0.5626', '14'],
        ['SMOTE', '0.9935', '1.0000', '0.9967', '16'],
        ['Class Weights', '0.9939', '0.9988', '0.9964', '15'],
        ['Combined', '0.9943', '0.4190', '0.5896', '14'],
    ]

    for row_idx, row_data in enumerate(imb_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table6.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph('Table 6: Imbalance Handling Methods Comparison')

    imb_disc = doc.add_paragraph()
    imb_disc.add_run('SMOTE achieved the highest F1-score (0.9967) with perfect precision and 99.35% recall. However, given the finding that full_drain dominates the predictions, these results primarily reflect the model\'s ability to identify the account drainage pattern rather than general fraud detection capability.')

    # =========================================================================
    # DISCUSSION - NEW SECTION
    # =========================================================================
    doc.add_heading('5.6 Discussion: Implications for Real-World Application', level=2)

    disc_note = doc.add_paragraph()
    run = disc_note.add_run('[NEW SECTION - Critical discussion of findings]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    discussion = doc.add_paragraph()
    run = discussion.add_run('''The discovery that full_drain captures 97.5% of fraudulent transactions reveals an important characteristic of the PaySim dataset: fraud is defined by a specific, deterministic pattern (complete account drainage) rather than the diverse and sophisticated patterns observed in real-world fraud.

This has several implications for interpreting our results:

1. Methodological Validity: The machine learning pipeline, feature engineering approach, and evaluation methodology remain valid and demonstrate proper experimental design. The models correctly learned the patterns present in the data.

2. Metric Interpretation: The near-perfect metrics (AUC > 0.99, recall > 99%) should be interpreted as reflecting PaySim's simplified fraud definition rather than expected real-world performance. The realistic estimate (recall ~69% without full_drain) provides a more conservative baseline.

3. Feature Engineering Value: The full_drain feature demonstrates the power of domain-informed feature engineering—it successfully captures the core fraud pattern. In real-world applications, similar domain knowledge could be valuable, though fraud patterns would be more diverse.

4. Synthetic Data Limitations: This finding highlights a broader limitation of synthetic datasets: they may not capture the full complexity and variability of real-world fraud, leading to overly optimistic performance estimates.''')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # CONCLUSION
    # =========================================================================
    doc.add_heading('6. Conclusion & Future Work', level=1)
    doc.add_paragraph('(max 300 words, 5 marks) EMMA')

    conc = doc.add_paragraph()
    conc.add_run('This paper explored the application of machine learning to fraud detection in financial transactions using the PaySim synthetic financial dataset. The study addressed three research questions regarding algorithm selection, feature importance, and class imbalance handling.\n\n')

    run = conc.add_run('Key findings include: (1) Random Forest achieved the highest ROC-AUC (0.9992) with the full feature set; (2) The engineered full_drain feature dominated with 92.6% importance, capturing 97.5% of fraudulent transactions; (3) SMOTE achieved optimal precision-recall balance among imbalance handling methods.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    conc.add_run('\n\n')

    run = conc.add_run('Critically, our analysis revealed that PaySim\'s fraud follows a deterministic pattern—complete account drainage—which the full_drain feature essentially encodes. When this feature is excluded, model recall drops from 99% to 69%, providing a more realistic performance estimate for real-world applications where fraud patterns are more diverse.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    conc.add_run('\n\nThe project demonstrates both the power and limitations of synthetic datasets for fraud detection research. While PaySim enables development and validation of machine learning pipelines without privacy concerns, its simplified fraud definition limits generalizability of performance metrics.\n\n')

    conc.add_run('Future work should: (1) validate these methods on real transaction data where available; (2) investigate more sophisticated fraud patterns beyond account drainage; (3) explore graph-based detection capturing account relationships; and (4) develop adaptive learning methods for concept drift. Researchers using synthetic datasets should carefully analyze the relationship between engineered features and target definitions to avoid overly optimistic performance estimates.')

    # =========================================================================
    # ETHICS STATEMENT
    # =========================================================================
    doc.add_heading('7. Ethics Statement', level=1)
    doc.add_paragraph('SARAH')

    ethics = """Fraud detection systems raise important ethical concerns due to their impact on users and financial institutions. A key issue is the trade-off between false positives and false negatives. False positives may incorrectly block legitimate transactions, causing inconvenience, while false negatives allow fraudulent activity to go undetected. Bias and fairness are also important considerations, as models may learn patterns that unfairly target certain behaviours. Although this study uses the synthetic PaySim dataset, real-world systems must ensure fair and unbiased decision-making. Transparency is essential in financial applications, particularly when using complex models. This study addresses this through feature importance analysis to improve interpretability. Additionally, while synthetic data is used, real systems must ensure compliance with data privacy regulations such as GDPR."""
    doc.add_paragraph(ethics)

    # =========================================================================
    # AI DISCLOSURE
    # =========================================================================
    doc.add_heading('8. Use of Generative AI', level=1)

    ai_note = doc.add_paragraph()
    run = ai_note.add_run('[Required by assignment]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    ai_disclosure = doc.add_paragraph()
    run = ai_disclosure.add_run('Generative AI tools (Claude) were used to assist with code debugging, cross-validating results between the notebook and paper, identifying the data leakage issue with the full_drain feature, and suggesting improvements to the paper structure. All experimental results, analysis, and conclusions are based on the authors\' own work using the provided notebook code. The critical finding regarding full_drain\'s relationship with the fraud definition was identified through AI-assisted analysis and subsequently verified by the authors.')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('9. References', level=1)

    references = """[1] Lopez-Rojas, E.A., Elmir, A., & Axelsson, S. (2016). PaySim: A financial mobile money simulator for fraud detection. Proceedings of the 28th European Modeling and Simulation Symposium (EMSS), Larnaca, Cyprus.

[2] Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLOS ONE, 10(3), 1-21.

[3] Lopez-Rojas, E. (2016). Synthetic Financial Datasets For Fraud Detection. Kaggle.com.

[4] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.

[5] Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30.

[6] Chawla, N.V., Bowyer, K.W., Hall, L.O., & Kegelmeyer, W.P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321-357.

[7] Dal Pozzolo, A., Caelen, O., Johnson, R.A., & Bontempi, G. (2015). Calibrating probability with undersampling for unbalanced classification. IEEE Symposium Series on Computational Intelligence.

[8] Dornadula, V.N., & Geetha, S. (2019). Credit card fraud detection using machine learning algorithms. Procedia Computer Science, 165, 631-641.

[9] Sailusha, R., Gnaneswar, V., Ramesh, R., & Rao, G.R. (2020). Credit card fraud detection using machine learning. IEEE Xplore.

[10] Velarde, G., Sudhir, A., Deshmane, S., Deshmukh, A., Sharma, K., & Joshi, V. (2023). Evaluating XGBoost for balanced and imbalanced data: Application to fraud detection. arXiv preprint, Cornell University.

[11] Hajek, P., Abedin, M.Z., & Sivarajah, U. (2022). Fraud detection in mobile payment systems using an XGBoost-based framework. Information Systems Frontiers, 25.

[12] RB, A., & KR, S.K. (2021). Credit card fraud detection using artificial neural network. Global Transitions Proceedings, 2(1).

[13] Rubaidi, Z.S., Ammar, B.B., & Aouicha, M.B. (2022). Fraud detection using large-scale imbalance dataset. International Journal on Artificial Intelligence Tools.

[14] Adebayo, O.S., Favour-Bethy, T.A., Owolafe, O., & Okunola, O.A. (2023). Comparative review of credit card fraud detection using machine learning and concept drift techniques. International Journal of Computer Science and Mobile Computing, 12(7), 24-48.

[15] Lin, K., & Gao, Y. (2022). Model interpretability of financial fraud detection by group SHAP. Expert Systems with Applications, 210, 118354.

[16] Wang, H., Liang, Q., Hancock, J.T., & Khoshgoftaar, T.M. (2024). Feature selection strategies: A comparative analysis of SHAP-value and importance-based methods. Journal of Big Data, 11(1)."""

    doc.add_paragraph(references)

    # Save the document
    output_path = "D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - FINAL.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_updated_paper_v2()
