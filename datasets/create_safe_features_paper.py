"""
Create a new Word document with results from the SAFE features analysis.
This creates a copy - does not modify the original.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_paper():
    doc = Document()

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('Fraud Detection in Financial Transactions Using Machine Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('An Analysis of the PaySim Synthetic Financial Dataset')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Lee Hickey | Emma Reen | Sarah O\'Hanlon\n').bold = True
    authors.add_run('lee.hickey28@mail.dcu.ie | emma.reen2@mail.dcu.ie | sarah.ohanlon24@mail.dcu.ie\n')
    authors.add_run('21407166 | 21491492 | 22444924\n')
    authors.add_run('School of Computing, Faculty of Engineering and Computing\n')
    authors.add_run('Dublin City University')

    doc.add_paragraph()

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    doc.add_heading('Abstract', level=1)

    abstract = doc.add_paragraph()
    abstract.add_run(
        'Financial fraud poses a significant threat to the integrity of digital payment systems, '
        'resulting in substantial economic losses worldwide. This study investigates the application '
        'of machine learning techniques for detecting fraudulent transactions in mobile money systems. '
        'Using the PaySim synthetic financial dataset, which simulates real-world mobile money transactions, '
        'we develop and evaluate classification models capable of accurately identifying fraudulent activities. '
        'The dataset contains over 6.3 million transactions across five transaction types, with approximately '
        '8,213 labeled fraudulent cases representing a highly imbalanced classification problem (0.13% fraud rate). '
    )

    abstract.add_run(
        'Critically, our analysis identified and addressed a data leakage issue in commonly used features, '
        'ensuring our results reflect realistic deployment performance. '
    ).bold = True

    abstract.add_run(
        'Using only pre-transaction features (safe features), our results demonstrate that '
    )
    abstract.add_run('XGBoost achieves the highest ROC-AUC (0.999) ').bold = True
    abstract.add_run(
        'among tested algorithms, with the amount-to-balance ratio emerging as the most predictive feature (89.4% importance). '
        'For class imbalance handling, Weighted XGBoost achieves 96.6% recall while maintaining 46.2% precision, '
        'representing realistic performance expectations for production deployment.'
    )

    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    doc.add_heading('1. Introduction', level=1)

    intro_text = """The growth of digital financial services has transformed how individuals and businesses transfer money. Mobile banking, digital wallets, and online payment systems now process millions of transactions each day improving convenience, speed, and accessibility. However, this expansion has also increased the risk of financial fraud. Fraudulent transactions can cause financial losses, reduce customer trust, and create legal and reputational problems for the service providers. As a result, fraud detection has become an important application of data science and artificial intelligence in finance.

Traditional fraud detection systems have often relied on rule-based methods. These systems flag transactions based on fixed conditions, such as unusually large amounts or repeated transfers. Although this approach can detect simple and known fraud patterns, it is often ineffective against more complex or changing behaviour. Fraudsters can adapt their actions to avoid these predefined rules, which limits the effectiveness of static systems. Machine learning offers a more flexible alternative because it can learn patterns directly from historical transaction data. Instead of depending only on manually created rules, machine learning models can detect complex relationships between features and identify suspicious behaviour more effectively.

Despite its advantages, fraud detection remains a difficult machine learning task. One major challenge is class imbalance. In most transactions datasets, fraudulent cases account for only a very small proportion of all records. This means a classifier can achieve high overall accuracy by predicting nearly all transactions as legitimate, while still failing to identify the fraud cases that matter most. For this reason evaluation metrics such as precision, recall, F1-score, and ROC-AUC are more useful than accuracy alone when assessing fraud detection models. In real-world applications, false negatives are particularly costly as they allow fraudulent transactions to go undetected.

Another challenge is that fraudulent behaviour is often hidden within normal transaction activity. Raw features such as amount, transaction type, or account balance may not be sufficient on their own to identify suspicious behaviour. Therefore, feature engineering is often necessary to improve model performance. Features based on balance changes, account draining, transaction ratios, and deviations from normal account behaviour may reveal useful fraud patterns.

This paper investigates fraud detection using the PaySim synthetic financial dataset. PaySim was created to simulate mobile money transactions using aggregated characteristics from real financial logs, while avoiding the privacy issues associated with real customer data. The dataset contains 6,362,620 transactions and 11 original features, with only 8,213 cases labelled as fraudulent, representing approximately 0.13% of the total data. It includes five transaction types: CASH_IN, CASH_OUT, DEBIT, PAYMENT, and TRANSFER.

The main aim of this study is to evaluate how effectively machine learning algorithms can detect fraudulent financial transactions in a highly imbalanced dataset using only features available at transaction time. This paper addresses three research questions:

1. Which machine learning algorithms perform best for fraud detection in imbalanced transaction data?
2. Which features are most predictive of fraudulent activity?
3. How can class imbalance be handled effectively?"""

    doc.add_paragraph(intro_text)

    # =========================================================================
    # RELATED WORK
    # =========================================================================
    doc.add_heading('2. Related Work', level=1)

    related_work_text = """Financial fraud detection has become a widely studied problem due to the rapid growth of digital and mobile transactions. Machine learning techniques are increasingly used to identify fraudulent behaviour by analysing transaction patterns at scale. However, fraud detection presents unique challenges, particularly due to its temporal nature and the risk of unrealistic evaluation.

A significant body of research has focused on comparing machine learning algorithms for fraud detection. Studies demonstrate that ensemble methods, particularly Random Forest and gradient boosting techniques such as XGBoost, outperform simpler models due to their ability to capture complex, nonlinear relationships in transaction data. These findings align with this study, which evaluates a range of machine learning models to determine their effectiveness in detecting fraudulent transactions.

A fundamental challenge in fraud detection is the severe class imbalance present in financial transaction datasets, where fraudulent transactions typically represent a very small proportion of the data. This imbalance can significantly bias machine learning models towards the majority class, resulting in high overall accuracy but poor detection of fraudulent cases. Various techniques, including oversampling (SMOTE), undersampling, and cost-sensitive learning, have been proposed to address this issue.

Another critical aspect of fraud detection is the risk of data leakage - where features used for prediction contain information that would not be available at prediction time in a real deployment scenario. This study specifically addresses this concern by using only pre-transaction features for model training and evaluation."""

    doc.add_paragraph(related_work_text)

    # =========================================================================
    # METHODOLOGY
    # =========================================================================
    doc.add_heading('3. Methodology', level=1)

    doc.add_heading('3.1 Dataset Description', level=2)
    dataset_text = """The PaySim dataset simulates mobile money transactions based on real transaction logs. The dataset contains:
• 6,362,620 total transactions
• 8,213 fraudulent transactions (0.13%)
• 5 transaction types: CASH_IN, CASH_OUT, DEBIT, PAYMENT, TRANSFER
• Class imbalance ratio: approximately 773:1"""
    doc.add_paragraph(dataset_text)

    doc.add_heading('3.2 Feature Engineering - Safe Features', level=2)

    safe_features_text = """A critical aspect of this study is the identification and prevention of data leakage. During initial analysis, we discovered that certain engineered features (particularly 'full_drain' - indicating complete account drainage) had near-perfect correlation with the fraud label, essentially encoding the fraud definition itself.

To ensure realistic and deployable results, we restricted our analysis to SAFE features - those available BEFORE a transaction completes:"""
    doc.add_paragraph(safe_features_text)

    # Safe features table
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Table Grid'
    headers = ['Feature', 'Description']
    table.rows[0].cells[0].text = headers[0]
    table.rows[0].cells[1].text = headers[1]
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    features = [
        ('step', 'Time step of the transaction (hourly)'),
        ('type_encoded', 'Transaction type (encoded)'),
        ('amount', 'Transaction amount'),
        ('oldbalanceOrg', 'Sender balance before transaction'),
        ('oldbalanceDest', 'Recipient balance before transaction'),
        ('amount_to_orig_balance', 'Ratio of amount to sender balance'),
        ('amount_to_dest_balance', 'Ratio of amount to recipient balance'),
        ('dest_zero_balance_before', 'Flag if recipient has zero balance'),
        ('hour_of_day', 'Hour of day (0-23)'),
        ('day', 'Day of simulation')
    ]
    for i, (feat, desc) in enumerate(features, 1):
        table.rows[i].cells[0].text = feat
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 1: Safe Features Used for Fraud Detection')

    doc.add_heading('3.3 Machine Learning Algorithms', level=2)
    algorithms_text = """Five machine learning algorithms were evaluated:
• Logistic Regression - Linear baseline model
• Decision Tree - Single tree classifier (max_depth=10)
• Random Forest - Ensemble of 50 decision trees
• Gradient Boosting - Sequential ensemble method
• XGBoost - Optimized gradient boosting implementation"""
    doc.add_paragraph(algorithms_text)

    doc.add_heading('3.4 Class Imbalance Handling', level=2)
    imbalance_text = """Five approaches to handling class imbalance were evaluated:
• Baseline - No imbalance handling
• Random Undersampling - Reduce majority class samples
• SMOTE - Synthetic Minority Over-sampling Technique
• Class Weighting - Penalize misclassification of minority class
• Combined - Undersampling with class weighting"""
    doc.add_paragraph(imbalance_text)

    # =========================================================================
    # EXPERIMENTS
    # =========================================================================
    doc.add_heading('4. Experiments', level=1)

    experiments_text = """The dataset was split into 70% training and 30% test sets using stratified sampling to preserve class distribution. Due to computational constraints, a sample of 100,000 transactions was used for training and testing.

All models were trained using scikit-learn and XGBoost libraries in Python. Evaluation metrics included ROC-AUC, Precision, Recall, and F1-Score, with particular emphasis on Recall (fraud detection rate) given the business importance of catching fraudulent transactions."""
    doc.add_paragraph(experiments_text)

    # =========================================================================
    # RESULTS
    # =========================================================================
    doc.add_heading('5. Results', level=1)

    # RQ1 Results
    doc.add_heading('5.1 RQ1: Algorithm Comparison', level=2)

    rq1_intro = """Table 2 presents the performance comparison of machine learning algorithms using safe features (no data leakage)."""
    doc.add_paragraph(rq1_intro)

    # RQ1 Table
    table_rq1 = doc.add_table(rows=6, cols=5)
    table_rq1.style = 'Table Grid'
    rq1_headers = ['Algorithm', 'AUC', 'Precision', 'Recall', 'F1 Score']
    for i, h in enumerate(rq1_headers):
        table_rq1.rows[0].cells[i].text = h
        table_rq1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    rq1_data = [
        ('XGBoost', '0.9999', '0.936', '0.740', '0.826'),
        ('Random Forest', '0.9792', '0.959', '0.395', '0.560'),
        ('Gradient Boosting', '0.9558', '0.905', '0.319', '0.472'),
        ('Decision Tree', '0.8393', '0.787', '0.403', '0.533'),
        ('Logistic Regression', '0.7806', '0.278', '0.042', '0.073')
    ]
    for i, row_data in enumerate(rq1_data, 1):
        for j, val in enumerate(row_data):
            table_rq1.rows[i].cells[j].text = val

    doc.add_paragraph()
    doc.add_paragraph('Table 2: Algorithm Performance Comparison (Safe Features)')

    rq1_findings = """Key Findings:
• XGBoost significantly outperforms all other algorithms with AUC of 0.9999
• XGBoost achieves the highest recall (74.0%) and F1 score (0.826)
• Logistic Regression performs poorly, indicating fraud patterns are highly non-linear
• Tree-based ensemble methods consistently outperform single models"""
    doc.add_paragraph(rq1_findings)

    # RQ2 Results
    doc.add_heading('5.2 RQ2: Feature Importance', level=2)

    rq2_intro = """Table 3 presents the feature importance ranking from the XGBoost model trained on safe features."""
    doc.add_paragraph(rq2_intro)

    # RQ2 Table
    table_rq2 = doc.add_table(rows=11, cols=2)
    table_rq2.style = 'Table Grid'
    table_rq2.rows[0].cells[0].text = 'Feature'
    table_rq2.rows[0].cells[1].text = 'Importance'
    table_rq2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table_rq2.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    rq2_data = [
        ('amount_to_orig_balance', '0.8938'),
        ('amount', '0.0802'),
        ('type_encoded', '0.0087'),
        ('hour_of_day', '0.0044'),
        ('step', '0.0029'),
        ('amount_to_dest_balance', '0.0026'),
        ('oldbalanceDest', '0.0022'),
        ('day', '0.0022'),
        ('oldbalanceOrg', '0.0020'),
        ('dest_zero_balance_before', '0.0011')
    ]
    for i, (feat, imp) in enumerate(rq2_data, 1):
        table_rq2.rows[i].cells[0].text = feat
        table_rq2.rows[i].cells[1].text = imp

    doc.add_paragraph()
    doc.add_paragraph('Table 3: Feature Importance Ranking (Safe Features)')

    rq2_findings = """Key Findings:
• The ratio of transaction amount to sender's balance (amount_to_orig_balance) dominates with 89.4% importance
• This indicates that fraudsters attempt to transfer unusually large proportions of account balances
• Transaction amount itself is the second most important feature (8.0%)
• Transaction type provides modest predictive value (0.9%)
• Temporal features (hour, day) have minimal importance"""
    doc.add_paragraph(rq2_findings)

    # RQ3 Results
    doc.add_heading('5.3 RQ3: Class Imbalance Handling', level=2)

    rq3_intro = """Table 4 presents the comparison of class imbalance handling methods using XGBoost with safe features."""
    doc.add_paragraph(rq3_intro)

    # RQ3 Table
    table_rq3 = doc.add_table(rows=6, cols=6)
    table_rq3.style = 'Table Grid'
    rq3_headers = ['Method', 'AUC', 'Precision', 'Recall', 'F1', 'Frauds Caught']
    for i, h in enumerate(rq3_headers):
        table_rq3.rows[0].cells[i].text = h
        table_rq3.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    rq3_data = [
        ('Baseline', '0.9999', '0.936', '0.740', '0.826', '88/119'),
        ('Undersampled', '0.9943', '0.036', '0.992', '0.070', '118/119'),
        ('SMOTE', '0.9996', '0.292', '0.958', '0.448', '114/119'),
        ('Weighted', '0.9998', '0.462', '0.966', '0.625', '115/119'),
        ('Combined', '0.9966', '0.054', '0.992', '0.103', '118/119')
    ]
    for i, row_data in enumerate(rq3_data, 1):
        for j, val in enumerate(row_data):
            table_rq3.rows[i].cells[j].text = val

    doc.add_paragraph()
    doc.add_paragraph('Table 4: Class Imbalance Handling Comparison (Safe Features)')

    rq3_findings = """Key Findings:
• Weighted XGBoost achieves the best balance: 96.6% recall with 46.2% precision
• Undersampling achieves highest recall (99.2%) but very low precision (3.6%)
• SMOTE provides good recall (95.8%) with moderate precision (29.2%)
• Baseline has highest precision but misses 26% of fraudulent transactions
• The choice of method depends on business requirements - cost of missed fraud vs. false alarms"""
    doc.add_paragraph(rq3_findings)

    # =========================================================================
    # DISCUSSION
    # =========================================================================
    doc.add_heading('6. Discussion', level=1)

    doc.add_heading('6.1 Data Leakage Prevention', level=2)
    leakage_text = """A critical contribution of this study is the identification and prevention of data leakage. Initial experiments using all engineered features achieved near-perfect metrics (>99% across all measures). However, analysis revealed that the 'full_drain' feature - indicating complete account drainage - had a +0.987 correlation with fraud, essentially encoding the fraud definition.

By restricting analysis to safe features (those available before transaction completion), our results represent realistic deployment performance rather than artificially inflated metrics."""
    doc.add_paragraph(leakage_text)

    doc.add_heading('6.2 Practical Implications', level=2)
    practical_text = """For practitioners deploying fraud detection systems:

1. Algorithm Selection: XGBoost is strongly recommended due to its superior performance on safe features
2. Feature Engineering: Focus on transaction-to-balance ratios as primary fraud indicators
3. Imbalance Handling: Use class weighting for best precision-recall balance
4. Threshold Tuning: Consider lowering decision threshold (0.3-0.4) to catch more fraud at the cost of more false positives
5. Expectations: Real-world performance will be closer to our safe feature results (~74-97% recall) rather than inflated metrics from leaked features"""
    doc.add_paragraph(practical_text)

    # =========================================================================
    # CONCLUSION
    # =========================================================================
    doc.add_heading('7. Conclusion', level=1)

    conclusion_text = """This study demonstrates both the promise and challenges of machine learning for fraud detection in financial transactions. Our analysis of the PaySim dataset yielded several important findings:

Research Question 1: XGBoost achieved the highest ROC-AUC (0.9999) among tested algorithms using safe features, with tree-based ensemble methods consistently outperforming linear models.

Research Question 2: The ratio of transaction amount to sender's balance emerged as the dominant predictive feature (89.4% importance), indicating that fraudsters attempt to transfer unusually large proportions of account balances.

Research Question 3: Weighted XGBoost provided the best balance for handling class imbalance, achieving 96.6% recall while maintaining 46.2% precision.

Critical Finding: By identifying and preventing data leakage, this study provides realistic performance benchmarks for fraud detection deployment:
• Realistic AUC: 0.9999
• Realistic Recall: 74-97% (depending on imbalance handling)
• Realistic Precision: 29-94% (precision-recall trade-off)

These findings have important implications for practitioners. While near-perfect metrics may appear in research settings using leaked features, production systems should expect performance closer to our safe feature results. Organizations should plan for the operational impact of false positives while maintaining high fraud catch rates.

Future work should explore temporal validation approaches, concept drift detection, and real-time feature computation strategies for deployment scenarios."""
    doc.add_paragraph(conclusion_text)

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)

    references = [
        "[1] Lopez-Rojas, E.A., Elmir, A., Axelsson, S. (2016). PaySim: A financial mobile money simulator for fraud detection.",
        "[2] Chawla, N.V. et al. (2002). SMOTE: Synthetic Minority Over-sampling Technique.",
        "[3] Chen, T., Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System.",
        "[4] Dal Pozzolo, A. et al. (2015). Credit Card Fraud Detection: A Realistic Modeling and a Novel Learning Strategy.",
        "[5] Breiman, L. (2001). Random Forests. Machine Learning."
    ]
    for ref in references:
        doc.add_paragraph(ref)

    # Save document
    output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - SAFE FEATURES.docx'
    doc.save(output_path)

    print("="*60)
    print("PAPER CREATED SUCCESSFULLY")
    print("="*60)
    print(f"\nSaved to: {output_path}")
    print("\nContents:")
    print("  - Abstract (updated with safe features results)")
    print("  - Introduction")
    print("  - Related Work")
    print("  - Methodology (including safe features explanation)")
    print("  - Experiments")
    print("  - Results:")
    print("    - RQ1: Algorithm Comparison (Table 2)")
    print("    - RQ2: Feature Importance (Table 3)")
    print("    - RQ3: Imbalance Handling (Table 4)")
    print("  - Discussion (data leakage prevention)")
    print("  - Conclusion")
    print("  - References")
    print("="*60)

if __name__ == '__main__':
    create_paper()
