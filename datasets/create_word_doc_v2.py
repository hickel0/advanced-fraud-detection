from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(12)
    p.space_before = Pt(18 if level == 1 else 14)

def add_para(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Title Page
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Fraud Detection in Financial Transactions Using Machine Learning')
run.bold = True
run.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('An Analysis of the PaySim Synthetic Financial Dataset')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run('Authors\n').bold = True
authors.add_run('Lee Hickey\nEmma Reen\nSarah O\'Hanlon')

doc.add_paragraph()
abstract = doc.add_paragraph()
abstract.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract.add_run('Abstract\n').bold = True

abstract_text = """Financial fraud poses a significant threat to digital payment systems worldwide. This study investigates machine learning techniques for detecting fraudulent transactions using the PaySim synthetic dataset containing 6.3 million transactions. We compare multiple algorithms including Logistic Regression, Decision Trees, Random Forest, Gradient Boosting, XGBoost, LightGBM, and Neural Networks (TensorFlow/Keras). Feature engineering captures behavioral transaction patterns through balance dynamics, account behavior indicators, and temporal features. SHAP analysis provides model interpretability, while concept drift detection using rolling window evaluation and Page-Hinkley tests monitors performance degradation. Results demonstrate that tree-based ensemble methods achieve superior performance, with SMOTE and class weighting effectively addressing the 770:1 class imbalance."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(abstract_text)
run.italic = True
run.font.size = Pt(11)
p.paragraph_format.left_indent = Cm(1.27)
p.paragraph_format.right_indent = Cm(1.27)

doc.add_page_break()

# 1. Introduction
add_heading('1. Introduction')
add_para("""The rapid digitalization of financial services has revolutionized monetary transactions globally. Mobile money platforms, digital wallets, and online banking have become integral to the financial ecosystem, enabling seamless peer-to-peer transfers and merchant payments. However, these advances have created new opportunities for fraudulent activities, with global losses exceeding billions annually.""")

add_para("""Traditional fraud detection methods relying on rule-based systems have proven inadequate against sophisticated modern fraud tactics. These legacy systems operate on predefined thresholds, making them vulnerable to novel attack vectors. Furthermore, the sheer volume of transactions processed by contemporary platforms renders manual review impractical, necessitating automated machine learning-based detection mechanisms.""")

add_para("""Machine learning addresses these limitations by learning complex patterns from historical data, identifying subtle fraud indicators that escape rule-based systems. These algorithms adapt to evolving fraud patterns, continuously improving detection capabilities. However, fraud detection presents unique challenges, particularly the extreme class imbalance where fraudulent transactions typically constitute less than 1% of total volume.""")

add_para("""This study investigates machine learning techniques using the PaySim synthetic financial dataset, which simulates real-world mobile money transactions from an African service provider. The dataset comprises over 6.3 million transactions with approximately 8,213 labeled fraudulent cases (0.13% fraud rate), representing a highly imbalanced classification problem.""")

add_para("""Our research addresses three fundamental questions: (1) Which machine learning algorithms perform best for detecting fraud in imbalanced financial data? (2) What transaction features are most predictive of fraudulent behavior? (3) How can class imbalance be effectively addressed to minimize false negatives while maintaining acceptable precision? Additionally, we implement explainable AI using SHAP and concept drift detection using rolling window evaluation.""")

add_para("""The contributions include: comprehensive algorithm comparison including neural networks and LightGBM; behavioral feature engineering capturing account-level transaction patterns; SHAP-based model interpretability analysis; and concept drift detection methodologies for production monitoring.""")

# 2. Related Work
add_heading('2. Related Work')
add_heading('2.1 Machine Learning for Fraud Detection', 2)
add_para("""Early fraud detection relied on rule-based approaches with expert-defined thresholds. Bolton and Hand (2002) highlighted limitations of these systems in adapting to evolving fraud patterns. The transition to machine learning brought decision trees and ensemble methods. Bhattacharyya et al. (2011) demonstrated that random forests outperformed individual decision trees and neural networks for credit card fraud.""")

add_para("""Gradient boosting methods have emerged as state-of-the-art for tabular data classification. Chen and Guestrin (2016) introduced XGBoost, achieving exceptional performance through regularization and efficient parallel computation. Ke et al. (2017) proposed LightGBM, offering faster training through histogram-based algorithms and lower memory usage while maintaining comparable accuracy.""")

add_heading('2.2 Deep Learning Approaches', 2)
add_para("""Neural networks have been explored for fraud detection, though advantages over gradient boosting on tabular data remain debated. Fiore et al. (2019) proposed generative adversarial networks for generating synthetic fraud samples. However, computational requirements and reduced interpretability have limited deep learning adoption in production systems where model explainability is often required.""")

add_heading('2.3 Handling Class Imbalance', 2)
add_para("""Extreme class imbalance has motivated extensive research into resampling techniques. Chawla et al. (2002) introduced SMOTE, generating synthetic minority samples through interpolation. Liu et al. (2009) proposed EasyEnsemble, combining undersampling with ensemble learning. Cost-sensitive learning assigns higher misclassification costs to minority class errors, with implementations like XGBoost's scale_pos_weight parameter.""")

add_heading('2.4 Explainable AI and Concept Drift', 2)
add_para("""Model interpretability has gained importance for regulatory compliance. Lundberg and Lee (2017) introduced SHAP (SHapley Additive exPlanations), providing both global and local interpretability through game-theoretic feature attribution. Concept drift detection monitors model performance degradation as fraud patterns evolve. Bifet and Gavalda (2007) proposed ADWIN for adaptive windowing, while the Page-Hinkley test detects changes in sequential data means.""")

# 3. Methodology
add_heading('3. Methodology')
add_heading('3.1 Dataset Description', 2)
add_para("""The PaySim synthetic dataset comprises 6,362,620 transactions simulating one month of mobile money activity. Features include step (hourly time unit), type (CASH_IN, CASH_OUT, DEBIT, PAYMENT, TRANSFER), amount, account identifiers, and balance information before/after transactions. The target variable isFraud contains 8,213 positive cases (0.13%), representing a 770:1 imbalance ratio.""")

add_heading('3.2 Behavioral Feature Engineering', 2)
add_para("""We engineered features capturing behavioral transaction patterns across six categories:""")

add_para("""Balance Dynamics: Balance change features (newbalance - oldbalance), balance ratios, and error flags detecting discrepancies between expected and actual post-transaction balances indicating potential manipulation.""")

add_para("""Account Behavior Indicators: Full account drain detection where transaction amount equals original balance, near-drain flags (>90% of balance), and zero-balance indicators capturing complete fund withdrawal patterns.""")

add_para("""Transaction Size Features: Amount-to-balance ratios normalizing transactions relative to account capacity, log-transformed amounts handling skewness, and large transaction flags.""")

add_para("""Temporal Patterns: Hour-of-day (step mod 24), day-of-simulation, night transaction flags (10PM-5AM), and weekend indicators capturing time-based fraud patterns.""")

add_para("""Behavioral Aggregations: Per-account transaction frequency, average transaction amounts, deviation from account averages, and significant deviation flags (>2x average) capturing unusual activity relative to historical behavior.""")

add_heading('3.3 Machine Learning Models', 2)
add_para("""We compared seven classification algorithms: Logistic Regression (baseline linear model), Decision Tree, Random Forest (bagging ensemble), Gradient Boosting, XGBoost, LightGBM, and Neural Network. LightGBM was included for its computational efficiency through histogram-based learning and leaf-wise tree growth.""")

add_para("""The Neural Network architecture comprised: Input layer accepting all features; two hidden layers (128, 64 neurons) with ReLU activation, Batch Normalization, and 30% Dropout; additional hidden layers (32, 16 neurons); and sigmoid output. Training used Adam optimizer, binary cross-entropy loss, class weights for imbalance handling, and early stopping on validation AUC.""")

add_heading('3.4 Explainable AI (SHAP)', 2)
add_para("""SHAP (SHapley Additive exPlanations) was employed for model interpretability. TreeExplainer computed SHAP values for XGBoost predictions, decomposing outputs into individual feature contributions. Summary plots visualized global feature importance; waterfall plots explained individual predictions; dependence plots revealed feature-value relationships and interactions.""")

add_heading('3.5 Concept Drift Detection', 2)
add_para("""Concept drift was monitored using rolling window evaluation. The model was trained on temporally-ordered first 50% of data, then evaluated using sliding windows (50,000 transactions, 25,000 step size) over remaining data. Page-Hinkley test detected significant deviations in rolling AUC values, identifying points where model performance degraded beyond historical thresholds.""")

# 4. Experiments
add_heading('4. Experiments')
add_heading('4.1 Algorithm Comparison (RQ1)', 2)
add_para("""Seven algorithms were trained on 500,000 stratified samples and evaluated on the 30% test set. All models used StandardScaler preprocessing. XGBoost and LightGBM incorporated scale_pos_weight equal to the imbalance ratio. The neural network used class weights and early stopping. Evaluation metrics included ROC-AUC, precision, recall, F1-score, and average precision.""")

add_heading('4.2 Feature Analysis (RQ2)', 2)
add_para("""XGBoost (200 estimators, depth 6) was trained for feature importance extraction. Built-in gain-based importance ranked features. SHAP TreeExplainer computed values for 3,000 test samples. Summary plots, bar plots, and waterfall plots visualized feature contributions. Correlation analysis computed Pearson coefficients with the target variable.""")

add_heading('4.3 Imbalance Handling (RQ3)', 2)
add_para("""Five approaches were compared using LightGBM: baseline (no handling), random undersampling, SMOTE oversampling (k=5), class weighting (scale_pos_weight), and combined (undersampling to 2:1 ratio with weight=2). Metrics focused on recall (fraud detection rate), precision, F1-score, and false negatives (missed frauds).""")

add_heading('4.4 Concept Drift Detection', 2)
add_para("""Rolling window evaluation computed AUC, precision, and recall for each 50,000-transaction window. Page-Hinkley test (delta=0.01, threshold=5) monitored AUC sequences for drift. Performance statistics (mean, std, min, max AUC) summarized temporal stability. Fraud rate monitoring identified data distribution shifts.""")

# 5. Results
add_heading('5. Results')
add_heading('5.1 Algorithm Performance (RQ1)', 2)
add_para("""Tree-based ensemble methods consistently outperformed other approaches. Random Forest and LightGBM achieved the highest AUC scores, with XGBoost close behind. LightGBM provided comparable performance to XGBoost with faster training times, confirming its suitability for production environments.""")

add_para("""The Neural Network achieved competitive performance but did not surpass gradient boosting methods on this tabular dataset. This aligns with literature suggesting that deep learning advantages are less pronounced for structured data compared to unstructured domains (images, text).""")

add_para("""All ensemble methods achieved recall above 0.99, correctly identifying over 99% of fraudulent transactions. Precision varied based on class imbalance handling, with SMOTE and class weighting providing the best precision-recall balance.""")

add_heading('5.2 Feature Importance (RQ2)', 2)
add_para("""XGBoost feature importance identified full_drain as the most predictive feature (importance > 0.9), followed by amount_to_dest_balance and balance-related features. This confirms domain knowledge that complete account drainage is strongly associated with fraud.""")

add_para("""SHAP analysis revealed that high transaction amounts, zero post-transaction balances, and TRANSFER/CASH_OUT types pushed predictions toward fraud. Behavioral aggregation features (amount deviation, significantly above average) captured unusual account activity patterns. The waterfall plots demonstrated how individual fraudulent transactions exhibited multiple suspicious indicators combining to produce high fraud probability predictions.""")

add_heading('5.3 Imbalance Handling (RQ3)', 2)
add_para("""SMOTE achieved the best F1-score, balancing precision and recall effectively through synthetic minority sample generation. Class weighting provided comparable results with lower computational overhead, making it preferable for large-scale production systems. Undersampling improved recall but sacrificed precision due to majority class information loss. The combined approach offered intermediate performance.""")

add_heading('5.4 Concept Drift Detection', 2)
add_para("""Rolling window evaluation showed mean AUC remaining stable above 0.99 throughout the evaluation period, indicating robust model generalization. However, periodic fluctuations in AUC (standard deviation ~0.01) and fraud rate variations demonstrated the need for continuous monitoring. Page-Hinkley test detected several drift points where performance temporarily deviated from historical averages, suggesting potential concept drift events requiring model adaptation.""")

# 6. Conclusion
add_heading('6. Conclusion and Future Work')
add_para("""This study investigated machine learning approaches for financial fraud detection, addressing algorithm selection, feature importance, and class imbalance handling, while implementing explainable AI and concept drift detection.""")

add_para("""Key findings: (1) Tree-based ensemble methods (XGBoost, LightGBM, Random Forest) achieve superior fraud detection performance compared to linear models and neural networks on tabular transaction data. (2) Behavioral features, particularly full account drainage and transaction-to-balance ratios, are most predictive of fraud. SHAP analysis confirms these patterns while providing interpretable explanations for individual predictions. (3) SMOTE and class weighting effectively address class imbalance, reducing false negatives while maintaining precision. (4) Concept drift monitoring reveals the need for continuous performance tracking and periodic model retraining.""")

add_para("""Future work should explore real-time streaming detection, graph neural networks capturing account relationship patterns, federated learning for privacy-preserving cross-institution collaboration, and automated concept drift adaptation through online learning approaches.""")

# 7. Ethics
add_heading('7. Ethics Statement')
add_para("""This research utilizes synthetic data from the PaySim simulator, preserving statistical properties of real transactions without exposing actual customer information. No real personal or financial data was accessed. We acknowledge dual-use concerns; our findings aim to improve security rather than inform adversarial attacks. Deployment of fraud detection systems requires fairness audits to prevent discrimination, and human oversight remains essential for high-stakes decisions.""")

# 8. References
add_heading('8. References')
refs = [
    "Bifet, A. & Gavalda, R. (2007). Learning from Time-Changing Data with Adaptive Windowing. SIAM SDM.",
    "Bhattacharyya, S. et al. (2011). Data mining for credit card fraud: A comparative study. Decision Support Systems.",
    "Bolton, R.J. & Hand, D.J. (2002). Statistical fraud detection: A review. Statistical Science.",
    "Chawla, N.V. et al. (2002). SMOTE: Synthetic Minority Over-sampling Technique. JAIR.",
    "Chen, T. & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. KDD.",
    "Fiore, U. et al. (2019). Using GANs for improving classification in credit card fraud detection. Information Sciences.",
    "Ke, G. et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. NeurIPS.",
    "Liu, X.Y. et al. (2009). Exploratory undersampling for class-imbalance learning. IEEE TSMC.",
    "Lopez-Rojas, E.A. et al. (2016). PaySim: A financial mobile money simulator for fraud detection. EMSS.",
    "Lundberg, S. & Lee, S.I. (2017). A Unified Approach to Interpreting Model Predictions (SHAP). NeurIPS."
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(ref)
    run.font.size = Pt(11)

doc.save(r'D:\Final Year\App Domains 3\datasets\Fraud_Detection_Report_v2.docx')
print("Document created: Fraud_Detection_Report_v2.docx")
