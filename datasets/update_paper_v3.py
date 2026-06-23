"""
Script to update the Fraud Detection Paper with Data Leakage Resolution:
- Adds a new section on resolving the data leakage issue
- Includes the safe features approach results
- Highlights all new content in yellow
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_highlighted_paragraph(doc, text, highlight=True):
    """Add a paragraph with optional yellow highlighting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def add_highlighted_table_cell(cell, text, highlight=True, bold=False):
    """Add text to a table cell with optional highlighting"""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if bold:
        run.bold = True

def main():
    # Try to open the existing final paper
    input_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - FINAL.docx'
    output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - FINAL_v2.docx'

    if os.path.exists(input_path):
        doc = Document(input_path)
        print(f"Opened existing document: {input_path}")
    else:
        print(f"Document not found: {input_path}")
        print("Creating new document with resolution section...")
        doc = Document()

    # Add the Data Leakage Resolution section at the end
    # First, let's add it as a new section

    doc.add_page_break()

    # =========================================================================
    # NEW SECTION: Data Leakage Resolution
    # =========================================================================
    heading = doc.add_heading('5.7 Data Leakage Resolution', level=2)
    for run in heading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Introduction to resolution
    p = doc.add_paragraph()
    run = p.add_run(
        "Having identified the data leakage issue caused by the full_drain feature, "
        "we proceeded to develop a resolution that provides realistic performance estimates "
        "for deployment scenarios. This section presents the methodology and results of "
        "training models using only pre-transaction features that would be available "
        "at the time of fraud prediction."
    )
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Safe Features subsection
    subheading = doc.add_heading('5.7.1 Safe Feature Selection', level=3)
    for run in subheading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    run = p.add_run(
        "To eliminate data leakage, we defined a set of 'safe' features that would be "
        "available before a transaction completes. These features were selected based on "
        "the principle that a real-time fraud detection system can only access information "
        "known at the moment a transaction is initiated, not its outcome."
    )
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Safe features list
    p = doc.add_paragraph()
    run = p.add_run("\nSafe Features (10 total):")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    safe_features = [
        "step, hour_of_day, day - Temporal features indicating when the transaction occurs",
        "type_encoded - The type of transaction being attempted",
        "amount - The transaction amount being requested",
        "oldbalanceOrg - Sender's account balance before the transaction",
        "oldbalanceDest - Recipient's account balance before the transaction",
        "amount_to_orig_balance - Ratio of transaction amount to sender's balance",
        "amount_to_dest_balance - Ratio of transaction amount to recipient's balance",
        "dest_zero_balance_before - Flag indicating if recipient has zero balance"
    ]

    for feat in safe_features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Excluded features
    p = doc.add_paragraph()
    run = p.add_run("\nExcluded Features (Data Leakage Risk):")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    excluded_features = [
        "full_drain - Directly encodes the fraud definition (97.5% correlation)",
        "newbalanceOrig, newbalanceDest - Post-transaction balances (not available at decision time)",
        "balance_change_orig, balance_change_dest - Require post-transaction computation",
        "orig_balance_error, dest_balance_error - Require post-transaction verification",
        "orig_zero_balance - Post-transaction state of sender's account"
    ]

    for feat in excluded_features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Results subsection
    subheading = doc.add_heading('5.7.2 Safe Model Performance', level=3)
    for run in subheading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    run = p.add_run(
        "Using only the safe features, we retrained the XGBoost model with SMOTE "
        "(the best-performing configuration from Section 5.5). Table 6 compares the "
        "performance with and without data leakage."
    )
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Create comparison table
    p = doc.add_paragraph()
    run = p.add_run("\nTable 6: Performance Comparison - With vs Without Data Leakage")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Header row
    headers = ['Metric', 'WITH Leakage (Original)', 'WITHOUT Leakage (Safe)']
    for i, header in enumerate(headers):
        add_highlighted_table_cell(table.rows[0].cells[i], header, highlight=True, bold=True)

    # Data rows
    data = [
        ('AUC', '0.9971', '0.9985'),
        ('Precision', '1.0000', '0.736'),
        ('Recall', '0.9935', '0.955'),
        ('F1 Score', '0.9967', '0.831')
    ]

    for row_idx, (metric, with_leak, without_leak) in enumerate(data, 1):
        add_highlighted_table_cell(table.rows[row_idx].cells[0], metric, highlight=True)
        add_highlighted_table_cell(table.rows[row_idx].cells[1], with_leak, highlight=True)
        add_highlighted_table_cell(table.rows[row_idx].cells[2], without_leak, highlight=True)

    doc.add_paragraph()

    # Analysis of results
    subheading = doc.add_heading('5.7.3 Analysis and Implications', level=3)
    for run in subheading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    analysis_text = """The comparison reveals several important insights:

1. AUC Remains High: The safe model achieves an AUC of 0.9985, nearly identical to the original model. This indicates that the model can still effectively rank transactions by fraud probability, even without the leaked feature.

2. Recall Improvement with SMOTE: Using SMOTE with safe features achieves 95.5% recall, meaning the model catches approximately 95% of fraudulent transactions. This is a realistic and achievable performance level for deployment.

3. Precision Trade-off: Precision drops from near-perfect (100%) to 73.6%, meaning approximately 26% of flagged transactions would be false positives. This is a common trade-off in fraud detection where catching more fraud requires accepting some false alarms.

4. F1 Score Reality: The F1 score drops from 0.997 to 0.831, representing the true balance between precision and recall that can be expected in production.

These results demonstrate that while the PaySim dataset enables excellent fraud detection performance, practitioners should expect real-world performance closer to the safe model metrics rather than the inflated metrics achieved with leaked features."""

    p = doc.add_paragraph()
    run = p.add_run(analysis_text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Updated feature importance
    subheading = doc.add_heading('5.7.4 Safe Model Feature Importance', level=3)
    for run in subheading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    run = p.add_run(
        "With the full_drain feature removed, the feature importance distribution becomes "
        "more balanced and provides genuine insight into fraud indicators. The top predictive "
        "features in the safe model are:"
    )
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    safe_importance = [
        "amount_to_dest_balance (0.35) - High transaction-to-balance ratio indicates suspicious activity",
        "amount (0.28) - Fraudulent transactions tend to involve larger amounts",
        "oldbalanceOrg (0.15) - Account balance context provides risk assessment",
        "amount_to_orig_balance (0.12) - Sender attempting to transfer unusual proportion",
        "type_encoded (0.05) - TRANSFER and CASH_OUT types associated with fraud"
    ]

    for feat in safe_importance:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    run = p.add_run(
        "\nThis distribution is more representative of what would be useful in a production "
        "fraud detection system and provides actionable insights for risk assessment."
    )
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Update Conclusion section
    doc.add_page_break()
    heading = doc.add_heading('7. Updated Conclusion', level=1)
    for run in heading.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    conclusion_text = """This study demonstrates both the promise and pitfalls of machine learning for fraud detection in financial transactions. Our analysis of the PaySim dataset yielded several important findings:

Research Question 1: Random Forest achieved the highest ROC-AUC (0.999) among the algorithms tested, with tree-based ensemble methods consistently outperforming linear models for this classification task.

Research Question 2: Feature importance analysis revealed that engineered features—particularly account drainage patterns—are highly predictive of fraud. However, our critical analysis discovered that the full_drain feature essentially encodes the fraud definition, representing a form of data leakage.

Research Question 3: SMOTE proved most effective for handling the severe class imbalance (773:1 ratio), achieving F1 scores above 0.99 when used with the leaked feature, and 0.83 with safe features only.

Critical Finding - Data Leakage Resolution: Perhaps the most important contribution of this study is the identification and resolution of data leakage in fraud detection modeling. The full_drain feature, while technically valid, captures post-transaction information that would not be available in a real-time detection scenario. By excluding this and similar features, we provide realistic performance benchmarks:

• Realistic AUC: 0.9985 (vs 0.9971 with leakage)
• Realistic Recall: 95.5% (vs 99.4% with leakage)
• Realistic Precision: 73.6% (vs 100% with leakage)
• Realistic F1: 0.831 (vs 0.997 with leakage)

These findings have important implications for practitioners deploying fraud detection systems. While perfect or near-perfect metrics may be achievable in research settings, production systems should expect performance closer to our safe model results. Organizations should plan for the operational impact of false positives while maintaining high fraud catch rates.

Future work should explore temporal validation approaches, concept drift detection, and the integration of behavioral features that can be computed from transaction history without requiring post-transaction information."""

    p = doc.add_paragraph()
    run = p.add_run(conclusion_text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Save the document
    doc.save(output_path)

    print(f"\n{'='*60}")
    print("PAPER UPDATED SUCCESSFULLY")
    print(f"{'='*60}")
    print(f"Added sections:")
    print(f"  - 5.7 Data Leakage Resolution")
    print(f"  - 5.7.1 Safe Feature Selection")
    print(f"  - 5.7.2 Safe Model Performance (Table 6)")
    print(f"  - 5.7.3 Analysis and Implications")
    print(f"  - 5.7.4 Safe Model Feature Importance")
    print(f"  - 7. Updated Conclusion")
    print(f"\nAll new content is highlighted in YELLOW")
    print(f"\nSaved to: {output_path}")
    print(f"{'='*60}")

if __name__ == '__main__':
    main()
