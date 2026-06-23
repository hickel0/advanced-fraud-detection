"""
Update the paper to include comprehensive references matching the original.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the document
doc_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - FINAL_v3.docx'
doc = Document(doc_path)

# Find and remove the existing References section
ref_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'References':
        ref_idx = i
        break

# Remove existing references (paragraphs after "References" heading)
if ref_idx is not None:
    # Count paragraphs to remove (references are after the heading)
    paras_to_remove = []
    for i in range(ref_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith('['):
            paras_to_remove.append(i)

    # Remove in reverse order
    for i in reversed(paras_to_remove):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

# Add comprehensive references
references = [
    "[1] Lopez-Rojas, E.A., Elmir, A., & Axelsson, S. (2016). PaySim: A financial mobile money simulator for fraud detection. In 28th European Modeling and Simulation Symposium (EMSS), pp. 249-255.",

    "[2] Chawla, N.V., Bowyer, K.W., Hall, L.O., & Kegelmeyer, W.P. (2002). SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321-357.",

    "[3] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785-794.",

    "[4] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",

    "[5] Dal Pozzolo, A., Caelen, O., Le Borgne, Y.A., Waterschoot, S., & Bontempi, G. (2014). Learned lessons in credit card fraud detection from a practitioner perspective. Expert Systems with Applications, 41(10), 4915-4928.",

    "[6] He, H., & Garcia, E.A. (2009). Learning from Imbalanced Data. IEEE Transactions on Knowledge and Data Engineering, 21(9), 1263-1284.",

    "[7] Dal Pozzolo, A., Boracchi, G., Caelen, O., Alippi, C., & Bontempi, G. (2018). Credit Card Fraud Detection: A Realistic Modeling and a Novel Learning Strategy. IEEE Transactions on Neural Networks and Learning Systems, 29(8), 3784-3797.",

    "[8] Dornadula, V.N., & Geetha, S. (2019). Credit Card Fraud Detection using Machine Learning Algorithms. Procedia Computer Science, 165, 631-641.",

    "[9] Sailusha, R., Gnaneswar, V., Ramesh, R., & Rao, G.R. (2020). Credit Card Fraud Detection Using Machine Learning. In 2020 4th International Conference on Intelligent Computing and Control Systems (ICICCS), pp. 1264-1270.",

    "[10] Awoyemi, J.O., Adetunmbi, A.O., & Oluwadare, S.A. (2017). Credit card fraud detection using machine learning techniques: A comparative analysis. In 2017 International Conference on Computing Networking and Informatics (ICCNI), pp. 1-9.",

    "[11] Bhattacharyya, S., Jha, S., Tharakunnel, K., & Westland, J.C. (2011). Data mining for credit card fraud: A comparative study. Decision Support Systems, 50(3), 602-613.",

    "[12] Phua, C., Lee, V., Smith, K., & Gayler, R. (2010). A Comprehensive Survey of Data Mining-based Fraud Detection Research. arXiv preprint arXiv:1009.6119.",

    "[13] Rubaidi, A., Mokhtar, S.A., & Sidi, F. (2020). A Systematic Literature Review on Credit Card Fraud Detection. In 2020 8th International Conference on Information Technology and Multimedia (ICIMU), pp. 364-369.",

    "[14] Prusti, D., & Rath, S.K. (2019). Fraudulent Transaction Detection in Credit Card by Applying Ensemble Machine Learning Techniques. In 2019 10th International Conference on Computing, Communication and Networking Technologies (ICCCNT), pp. 1-6.",

    "[15] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",

    "[16] Lemaitre, G., Nogueira, F., & Aridas, C.K. (2017). Imbalanced-learn: A Python Toolbox to Tackle the Curse of Imbalanced Datasets in Machine Learning. Journal of Machine Learning Research, 18(17), 1-5."
]

# Find the References heading and add references after it
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'References':
        # Add references after this paragraph
        # We need to insert after the heading
        break

# Add each reference as a new paragraph at the end
for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(-36)  # Hanging indent
    p.paragraph_format.left_indent = Pt(36)

# Save
output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - FINAL_v3.docx'
doc.save(output_path)

print("="*70)
print("REFERENCES UPDATED")
print("="*70)
print(f"\nAdded {len(references)} references:")
for ref in references:
    print(f"  {ref[:70]}...")
print(f"\nSaved to: {output_path}")
print("="*70)
