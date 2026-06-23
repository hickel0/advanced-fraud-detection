"""
Fix the Fraud Detection Paper to align with notebook results.
Removes LightGBM and neural network references, adds missing Table 4.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Load the document
doc = Document('D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper (1).docx')

changes_made = []

def replace_in_paragraph(paragraph, old_text, new_text, change_desc):
    """Replace text in a paragraph if found."""
    if old_text in paragraph.text:
        # Store the original formatting
        inline = paragraph.runs
        for run in inline:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                changes_made.append(change_desc)
                return True
        # If not found in individual runs, try full paragraph
        if old_text in paragraph.text:
            full_text = paragraph.text
            new_full_text = full_text.replace(old_text, new_text)
            # Clear and rewrite
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
            changes_made.append(change_desc)
            return True
    return False

# Process all paragraphs
for para in doc.paragraphs:
    # Fix 1: Remove LightGBM and neural networks from algorithm list in Introduction
    old1 = "logistic regression, decision trees, random forests, gradient boosting, XGBoost, LightGBM, and neural networks"
    new1 = "Logistic Regression, Decision Trees, Random Forests, Gradient Boosting, and XGBoost"
    replace_in_paragraph(para, old1, new1,
        "Removed LightGBM and neural networks from algorithm list in Introduction")

    # Fix 2: Remove neural network comparison sentence in Related Work
    old2 = "While neural networks have also been explored for fraud detection, their performance on tabular data is often comparable rather than superior to ensemble methods, despite their ability to model highly nonlinear relationships [12]."
    new2 = ""
    if old2 in para.text:
        replace_in_paragraph(para, old2, new2,
            "Removed neural network comparison sentence in Related Work")

    # Fix 3: Update the sentence that mentions "including ensemble approaches and neural networks"
    old3 = "including ensemble approaches and neural networks"
    new3 = "including tree-based ensemble approaches"
    replace_in_paragraph(para, old3, new3,
        "Updated 'ensemble approaches and neural networks' to 'tree-based ensemble approaches'")

# Now let's go through the document more carefully to find and fix issues
print("="*70)
print("SCANNING DOCUMENT FOR ISSUES...")
print("="*70)

issues_found = []
para_num = 0

for para in doc.paragraphs:
    para_num += 1
    text_lower = para.text.lower()

    # Check for any remaining LightGBM references
    if 'lightgbm' in text_lower:
        issues_found.append(f"Para {para_num}: Still contains 'LightGBM': {para.text[:100]}...")

    # Check for neural network references (but not in references section)
    if 'neural network' in text_lower and 'reference' not in text_lower:
        # Skip if it's a citation like [12]
        if 'artificial neural network' in text_lower or '[12]' in para.text:
            # This is likely in references, skip
            pass
        elif 'graph neural' in text_lower:
            # This is in future work, which is OK
            pass
        else:
            issues_found.append(f"Para {para_num}: Contains 'neural network': {para.text[:100]}...")

print(f"\nIssues still remaining: {len(issues_found)}")
for issue in issues_found:
    print(f"  - {issue}")

# Manual fixes for remaining issues
for para in doc.paragraphs:
    # The related work section mentions neural networks in context of citing other work
    # We need to rewrite that sentence

    # Check for the "These findings align with this study" sentence
    if "These findings align with this study" in para.text and "neural networks" in para.text:
        old_text = para.text
        new_text = old_text.replace(
            "including ensemble approaches and neural networks, to determine their effectiveness in detecting fraudulent transactions",
            "including tree-based ensemble methods, to determine their effectiveness in detecting fraudulent transactions"
        )
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        changes_made.append("Fixed 'These findings align...' paragraph to remove neural networks reference")

# Save the corrected document
output_path = 'D:/Final Year/App Domains 3/datasets/CSC1112 - Fraud Detection Paper - CORRECTED.docx'
doc.save(output_path)

print("\n" + "="*70)
print("CHANGES MADE:")
print("="*70)
for i, change in enumerate(changes_made, 1):
    print(f"{i}. {change}")

print("\n" + "="*70)
print("ADDITIONAL ISSUES TO MANUALLY CHECK:")
print("="*70)
print("""
1. TABLE 4 (Feature Importance) - Needs to be added manually
   The paper references "Table 4 presents the top 10 features" but it's missing.

   Suggested Table 4 content:
   +----+----------------------------+------------+
   |Rank| Feature                    | Importance |
   +----+----------------------------+------------+
   | 1  | amount_to_orig_balance     | 0.9770     |
   | 2  | amount                     | 0.0170     |
   | 3  | type_encoded               | 0.0031     |
   | 4  | day                        | 0.0005     |
   | 5  | hour_of_day                | 0.0005     |
   | 6  | oldbalanceOrg              | 0.0005     |
   | 7  | dest_zero_balance_before   | 0.0004     |
   | 8  | amount_to_dest_balance     | 0.0004     |
   | 9  | oldbalanceDest             | 0.0004     |
   | 10 | step                       | 0.0003     |
   +----+----------------------------+------------+

2. Reference [12] - This reference is for neural networks:
   "RB, A., & KR, S.K. (2021). Credit card fraud detection using artificial
   neural network."

   This reference is NO LONGER CITED in the paper after removing neural network
   discussion. You may want to:
   - Remove reference [12] entirely, OR
   - Keep it but renumber references if needed

3. Reference [5] mentions LightGBM:
   "Ke, G., et al. (2017). LightGBM: A highly efficient gradient boosting..."

   This reference may no longer be needed since LightGBM is not discussed.
   Consider removing it.

4. In Related Work, the sentence starting "More recent studies further
   highlight..." references [10], [11] for XGBoost and LightGBM. Since we
   removed LightGBM, verify this sentence still makes sense.
""")

print(f"\nCorrected document saved to: {output_path}")
print("="*70)
